import json
from datetime import datetime
from pathlib import Path

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy.orm import Session

from ..models import Asset, ConfigDiff, ConfigSnapshot, Event, IngestionBatch, ReferenceRecord, ReportExport
from ..storage import export_path
from .event_intelligence import summarize_event_patterns


def _resolve_gc_component(parameter_key: str | None, parameter_label: str | None) -> str | None:
    text = f'{parameter_key or ""} {parameter_label or ""}'.lower().replace('_', ' ')
    if 'methane' in text or 'metano' in text or ' c1 ' in text:
        return 'Metano'
    if 'ethane' in text or 'etano' in text or ' c2 ' in text:
        return 'Etano'
    if 'propane' in text or 'propano' in text or ' c3 ' in text:
        return 'Propano'
    if 'ic4' in text or 'isobutane' in text or 'iso butane' in text or 'i butane' in text:
        return 'i-Butano'
    if 'nc4' in text or 'normal butane' in text or 'n butane' in text:
        return 'n-Butano'
    if 'ic5' in text or 'isopentane' in text or 'iso pentane' in text or 'i pentane' in text:
        return 'i-Pentano'
    if 'nc5' in text or 'normal pentane' in text or 'n pentane' in text:
        return 'n-Pentano'
    if 'hexane' in text or 'heptane' in text or 'c6+' in text:
        return 'Hexano+'
    if 'nitrogen' in text or 'nitrogeno' in text or ' n2' in text:
        return 'Nitrogênio'
    if 'carbon dioxide' in text or 'co2' in text:
        return 'CO2'
    return None


def _collect_gc_components(diff_records: list[ConfigDiff]) -> list[str]:
    components = []
    for record in diff_records:
        component = _resolve_gc_component(record.parameter_key, record.parameter_label)
        if component and component not in components:
            components.append(component)
    return components


def _load_expected_gc_components(session: Session) -> list[str]:
    record = (
        session.query(ReferenceRecord)
        .filter(
            ReferenceRecord.entity_type == 'metering_reference',
            ReferenceRecord.record_key == 'gc_component_profile',
        )
        .one_or_none()
    )
    if record is None:
        return []
    metadata = json.loads(record.metadata_json or '{}')
    components = metadata.get('components')
    if not isinstance(components, list):
        return []
    result: list[str] = []
    for component in components:
        resolved = _resolve_gc_component(str(component), str(component))
        if resolved and resolved not in result:
            result.append(resolved)
    return result


def _collect_snapshot_gc_components(snapshot: ConfigSnapshot | None) -> list[str]:
    if snapshot is None:
        return []
    components: list[str] = []
    for parameter in snapshot.parameters:
        component = _resolve_gc_component(parameter.parameter_key, parameter.parameter_label)
        if component and component not in components:
            components.append(component)
    return components


def _collect_diff_context(session: Session, left_snapshot_id: int, right_snapshot_id: int) -> dict:
    left_snapshot = session.get(ConfigSnapshot, left_snapshot_id)
    right_snapshot = session.get(ConfigSnapshot, right_snapshot_id)
    if left_snapshot is None or right_snapshot is None:
        raise ValueError('Snapshot not found for report generation.')
    asset = session.get(Asset, left_snapshot.asset_id)
    diff_records = (
        session.query(ConfigDiff)
        .filter(
            ConfigDiff.left_snapshot_id == left_snapshot_id,
            ConfigDiff.right_snapshot_id == right_snapshot_id,
        )
        .order_by(ConfigDiff.severity.desc(), ConfigDiff.parameter_label.asc())
        .all()
    )
    related_events = (
        session.query(Event)
        .filter(Event.asset_id == left_snapshot.asset_id)
        .order_by(Event.occurred_at.desc())
        .limit(12)
        .all()
    )
    intelligence = summarize_event_patterns(session, left_snapshot.asset_id)
    gc_expected_components = _load_expected_gc_components(session)
    gc_current_components = _collect_snapshot_gc_components(right_snapshot)
    gc_changed_components = _collect_gc_components(diff_records)
    return {
        'scope_type': 'diff',
        'scope_id': left_snapshot_id,
        'asset_name': asset.flow_computer_tag if asset else left_snapshot.device_name or 'Ativo sem nome',
        'left_snapshot': left_snapshot,
        'right_snapshot': right_snapshot,
        'diff_records': diff_records,
        'events': related_events,
        'intelligence': intelligence,
        'gc_components': gc_changed_components,
        'gc_current_components': gc_current_components,
        'gc_expected_components': gc_expected_components,
        'gc_missing_components': [component for component in gc_expected_components if component not in gc_current_components],
    }


def _collect_batch_context(session: Session, batch_id: int) -> dict:
    batch = session.get(IngestionBatch, batch_id)
    if batch is None:
        raise ValueError('Batch not found for report generation.')
    batch_files = sorted(batch.files, key=lambda item: item.original_name.lower())
    snapshots = (
        session.query(ConfigSnapshot)
        .filter(ConfigSnapshot.file.has(batch_id=batch_id))
        .order_by(ConfigSnapshot.snapshot_at.desc())
        .all()
    )
    events = (
        session.query(Event)
        .filter(Event.file.has(batch_id=batch_id))
        .order_by(Event.occurred_at.desc())
        .limit(20)
        .all()
    )
    asset_names = sorted({snapshot.asset.flow_computer_tag for snapshot in snapshots if snapshot.asset})
    gc_expected_components = _load_expected_gc_components(session)
    return {
        'scope_type': 'batch',
        'scope_id': batch_id,
        'asset_name': batch.source_name,
        'batch': batch,
        'batch_files': batch_files,
        'assets_found': asset_names,
        'left_snapshot': None,
        'right_snapshot': None,
        'diff_records': [],
        'events': events,
        'intelligence': summarize_event_patterns(session, batch_id=batch_id),
        'snapshots': snapshots,
        'gc_components': [],
        'gc_current_components': [],
        'gc_expected_components': gc_expected_components,
        'gc_missing_components': gc_expected_components,
    }


def _build_markdown(context: dict) -> str:
    lines = [
        '# Relatorio Tecnico SGMed Inspector',
        '',
        '## Resumo Executivo',
        '',
        f"- Escopo: `{context['asset_name']}`",
    ]
    if context['left_snapshot'] and context['right_snapshot']:
        lines.extend(
            [
                f"- Versão de comparação A: `{context['left_snapshot'].device_name}`",
                f"- Versão de comparação B: `{context['right_snapshot'].device_name}`",
                f"- Total de diferenças encontradas: `{len(context['diff_records'])}`",
                f"- Eventos avaliados: `{context['intelligence']['total_events']}`",
            ]
        )
    else:
        lines.extend(
            [
                f"- Arquivo carregado: `{context['asset_name']}`",
                f"- Importação analisada: `{context['scope_id']}`",
                f"- Arquivos encontrados dentro da importação: `{len(context.get('batch_files', []))}`",
                f"- Equipamentos encontrados nesta importação: `{len(context.get('assets_found', []))}`",
                f"- Versões de configuração nesta importação: `{len(context.get('snapshots', []))}`",
                f"- Eventos avaliados nesta importação: `{context['intelligence']['total_events']}`",
            ]
        )

    lines.extend(['', '## O Que Chama Atencao Primeiro', ''])
    if context['scope_type'] == 'batch':
        if context['assets_found']:
            lines.append(f"- Equipamentos encontrados: {', '.join(f'`{asset}`' for asset in context['assets_found'])}")
        else:
            lines.append('- Nenhum equipamento identificado neste arquivo.')
        recurring = context['intelligence']['recurring_patterns']
        if recurring:
            top_item = recurring[0]
            lines.append(f"- Evento mais recorrente: `{top_item['title']}` ({top_item['count']}x)")
        else:
            lines.append('- Nenhum evento recorrente detectado neste arquivo.')
    else:
        if context['gc_components']:
            lines.append(f"- Componentes de GC com mudança nesta comparação: {', '.join(f'`{component}`' for component in context['gc_components'])}")
        else:
            lines.append('- Nenhuma mudança de componente de GC apareceu nesta comparação.')
        if context.get('gc_current_components'):
            lines.append(f"- Componentes de GC lidos no arquivo carregado: {', '.join(f'`{component}`' for component in context['gc_current_components'])}")
        else:
            lines.append('- Nenhum componente de GC foi lido no arquivo carregado.')
        if context.get('gc_expected_components'):
            if context.get('gc_missing_components'):
                lines.append(f"- Componentes esperados que ainda não apareceram: {', '.join(f'`{component}`' for component in context['gc_missing_components'])}")
            else:
                lines.append('- Todos os componentes esperados de GC já apareceram nesta versão.')
    lines.extend(['', '## O Que Mudou', ''])
    if not context['diff_records']:
        lines.append('- Nenhuma diferenca catalogada para este escopo.')
    else:
        for record in context['diff_records']:
            lines.append(
                f"- [{record.severity}] `{record.parameter_label}`: `{record.left_value}` -> `{record.right_value}` ({record.category})"
            )

    lines.extend(['', '## Cromatografia / GC', ''])
    if context['scope_type'] == 'batch':
        lines.append('- Este relatório de importação não compara GC entre versões. Use uma comparação para ver desvio por componente.')
    else:
        if context.get('gc_components'):
            lines.append(f"- Componentes com mudança: {', '.join(f'`{component}`' for component in context['gc_components'])}")
        else:
            lines.append('- Componentes com mudança: nenhum componente de GC mudou nesta comparação.')
        if context.get('gc_current_components'):
            lines.append(f"- Componentes lidos no arquivo carregado: {', '.join(f'`{component}`' for component in context['gc_current_components'])}")
        else:
            lines.append('- Componentes lidos no arquivo carregado: nenhum componente de GC foi encontrado.')
        if context.get('gc_expected_components'):
            lines.append(f"- Perfil esperado: {', '.join(f'`{component}`' for component in context['gc_expected_components'])}")
            if context.get('gc_missing_components'):
                lines.append(f"- Ainda não apareceram nesta versão: {', '.join(f'`{component}`' for component in context['gc_missing_components'])}")
            else:
                lines.append('- Ainda não apareceram nesta versão: nenhum. Todos os componentes esperados foram encontrados.')
        else:
            lines.append('- Perfil esperado: nenhum perfil de GC configurado.')

    lines.extend(['', '## Leitura dos Eventos', ''])
    recurring = context['intelligence']['recurring_patterns']
    chattering = context['intelligence']['chattering_patterns']
    windows = context['intelligence']['operator_windows']
    if not recurring and not chattering and not windows:
        lines.append('- Nenhum padrao adicional detectado.')
    else:
        for item in recurring:
            lines.append(f"- Recorrencia: `{item['title']}` ({item['count']}x) - {item['detail']}")
        for item in chattering:
            lines.append(f"- Chattering: `{item['title']}` ({item['count']}x) - {item['detail']}")
        for item in windows:
            lines.append(f"- Janela de operador: `{item['title']}` - {item['detail']}")

    lines.extend(['', '## Eventos Recentes', ''])
    if not context['events']:
        lines.append('- Nenhum evento correlato encontrado.')
    else:
        for event in context['events']:
            timestamp = event.occurred_at.isoformat(sep=' ') if event.occurred_at else 'sem horario'
            lines.append(f'- `{timestamp}` {event.message}')

    lines.extend(
        [
            '',
            '## Proximas Acoes',
            '',
            '- Validar diferencas metrologicas antes de promover nova referencia oficial.',
            '- Confirmar se alarmes recorrentes representam degradacao persistente ou janela operacional esperada.',
        ]
    )
    return '\n'.join(lines)


def _write_pdf(target: Path, context: dict) -> None:
    styles = getSampleStyleSheet()
    document = SimpleDocTemplate(str(target), pagesize=A4, title='Relatorio Tecnico SGMed Inspector')
    story = [
        Paragraph('Relatorio Tecnico SGMed Inspector', styles['Title']),
        Spacer(1, 12),
        Paragraph(f"Escopo: {context['asset_name']}", styles['Heading2']),
        Spacer(1, 8),
        Paragraph('Resumo Executivo', styles['Heading2']),
        Spacer(1, 6),
        Paragraph(
            f"Mudancas encontradas: {len(context['diff_records'])} | Eventos avaliados: {context['intelligence']['total_events']}",
            styles['BodyText'],
        ),
        Spacer(1, 12),
    ]

    if context['diff_records']:
        table_data = [['Parametro', 'Esquerda', 'Direita', 'Categoria', 'Severidade']]
        for record in context['diff_records'][:20]:
            table_data.append([
                record.parameter_label,
                record.left_value or '-',
                record.right_value or '-',
                record.category,
                record.severity,
            ])
        table = Table(table_data, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#BB5C2A')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D7C9B9')),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FBF7F1')),
                ]
            )
        )
        story.append(Paragraph('O Que Mudou', styles['Heading2']))
        story.append(Spacer(1, 6))
        story.append(table)
        story.append(Spacer(1, 12))

    story.append(Paragraph('Cromatografia / GC', styles['Heading2']))
    story.append(Spacer(1, 6))
    if context['scope_type'] == 'batch':
        story.append(Paragraph('Este relatório de importação não compara GC entre versões. Use uma comparação para ver desvio por componente.', styles['BodyText']))
    else:
        changed = ', '.join(context.get('gc_components', [])) or 'nenhum componente com mudança'
        current = ', '.join(context.get('gc_current_components', [])) or 'nenhum componente lido'
        expected = ', '.join(context.get('gc_expected_components', [])) or 'nenhum perfil configurado'
        missing = ', '.join(context.get('gc_missing_components', [])) or 'nenhum componente ausente'
        story.append(Paragraph(f'Componentes com mudança: {changed}', styles['BodyText']))
        story.append(Spacer(1, 4))
        story.append(Paragraph(f'Componentes lidos no arquivo carregado: {current}', styles['BodyText']))
        story.append(Spacer(1, 4))
        story.append(Paragraph(f'Perfil esperado: {expected}', styles['BodyText']))
        story.append(Spacer(1, 4))
        story.append(Paragraph(f'Componentes esperados ainda ausentes: {missing}', styles['BodyText']))
    story.append(Spacer(1, 12))

    story.append(Paragraph('Leitura dos Eventos', styles['Heading2']))
    story.append(Spacer(1, 6))
    items = context['intelligence']['recurring_patterns'] + context['intelligence']['chattering_patterns']
    if not items:
        story.append(Paragraph('Nenhum padrao adicional detectado.', styles['BodyText']))
    else:
        for item in items:
            story.append(Paragraph(f"{item['title']} ({item['count']}x) - {item['detail']}", styles['BodyText']))
            story.append(Spacer(1, 4))

    document.build(story)


def _write_docx(target: Path, context: dict) -> None:
    document = Document()
    document.add_heading('Relatorio Tecnico SGMed Inspector', 0)
    document.add_paragraph(f"Escopo: {context['asset_name']}")
    document.add_heading('Resumo Executivo', level=1)
    document.add_paragraph(
        f"Mudancas encontradas: {len(context['diff_records'])} | Eventos avaliados: {context['intelligence']['total_events']}"
    )

    document.add_heading('O Que Mudou', level=1)
    if context['diff_records']:
        table = document.add_table(rows=1, cols=5)
        headers = ['Parametro', 'Esquerda', 'Direita', 'Categoria', 'Severidade']
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for record in context['diff_records'][:25]:
            row = table.add_row().cells
            row[0].text = record.parameter_label
            row[1].text = record.left_value or '-'
            row[2].text = record.right_value or '-'
            row[3].text = record.category
            row[4].text = record.severity
    else:
        document.add_paragraph('Nenhuma diferenca catalogada para este escopo.')

    document.add_heading('Cromatografia / GC', level=1)
    if context['scope_type'] == 'batch':
        document.add_paragraph('Este relatório de importação não compara GC entre versões. Use uma comparação para ver desvio por componente.')
    else:
        document.add_paragraph(
            f"Componentes com mudança: {', '.join(context.get('gc_components', [])) or 'nenhum componente com mudança'}"
        )
        document.add_paragraph(
            f"Componentes lidos no arquivo carregado: {', '.join(context.get('gc_current_components', [])) or 'nenhum componente lido'}"
        )
        document.add_paragraph(
            f"Perfil esperado: {', '.join(context.get('gc_expected_components', [])) or 'nenhum perfil configurado'}"
        )
        document.add_paragraph(
            f"Componentes esperados ainda ausentes: {', '.join(context.get('gc_missing_components', [])) or 'nenhum componente ausente'}"
        )

    document.add_heading('Leitura dos Eventos', level=1)
    items = context['intelligence']['recurring_patterns'] + context['intelligence']['chattering_patterns']
    if not items:
        document.add_paragraph('Nenhum padrao adicional detectado.')
    else:
        for item in items:
            document.add_paragraph(f"{item['title']} ({item['count']}x) - {item['detail']}", style='List Bullet')

    document.add_heading('Proximas Acoes', level=1)
    document.add_paragraph('Validar diferencas metrologicas antes de promover nova referencia oficial.', style='List Bullet')
    document.add_paragraph(
        'Confirmar se alarmes recorrentes representam degradacao persistente ou janela operacional esperada.',
        style='List Bullet',
    )
    document.save(target)


def create_report(
    session: Session,
    batch_id: int | None = None,
    left_snapshot_id: int | None = None,
    right_snapshot_id: int | None = None,
    export_format: str = 'markdown',
) -> tuple[ReportExport, str]:
    if left_snapshot_id and right_snapshot_id:
        context = _collect_diff_context(session, left_snapshot_id, right_snapshot_id)
    elif batch_id is not None:
        context = _collect_batch_context(session, batch_id)
    else:
        raise ValueError('Report request requires a batch or a diff pair.')

    content = _build_markdown(context)
    timestamp = datetime.utcnow().strftime('%Y%m%d%H%M%S')
    extension_map = {'markdown': 'md', 'pdf': 'pdf', 'docx': 'docx'}
    file_extension = extension_map.get(export_format)
    if file_extension is None:
        raise ValueError('Unsupported report format.')
    name = f"report-{context['scope_type']}-{context['scope_id']}-{timestamp}.{file_extension}"
    target = export_path(name)

    if export_format == 'markdown':
        target.write_text(content, encoding='utf-8')
    elif export_format == 'pdf':
        _write_pdf(target, context)
    else:
        _write_docx(target, context)

    export = ReportExport(scope_type=context['scope_type'], scope_id=context['scope_id'], format=export_format, file_path=str(target))
    session.add(export)
    session.commit()
    session.refresh(export)
    return export, content
