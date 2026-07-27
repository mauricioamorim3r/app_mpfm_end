from __future__ import annotations

import json
import hashlib
import math
import re
import sqlite3
import zipfile
from collections import Counter, defaultdict
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

import pandas as pd


TODAY = date.today()
REPORTING_PERIOD_START = date(TODAY.year, TODAY.month, 1)
REPORTING_PERIOD_END = date(TODAY.year, TODAY.month + 1, 1) - timedelta(days=1) if TODAY.month < 12 else date(TODAY.year, 12, 31)
APP_DIR = Path(__file__).resolve().parents[1]
ROOT = APP_DIR.parent
OUTPUT = APP_DIR / "src" / "data" / "dashboard-data.json"
CONFIG_PATH = APP_DIR / "config" / "data-sources.json"
SQLITE_PATH = APP_DIR / "data" / "radar-anp.sqlite"
CANDIDATE_DECISIONS_PATH = APP_DIR / "data" / "proposal-decisions.json"
PENDENCY_DECISIONS_PATH = APP_DIR / "data" / "pendency-decisions.json"
MODEL_DIR = APP_DIR / "MODELOS"
MAX_MODEL_FILE_AGGREGATES = 100_000
MAX_MODEL_OUTPUT_AGGREGATES = 5_000

FAMILIES = {
    "a001": {"name": "PMO", "label": "Oleo linear", "anp_file": "Óleo Linear.xlsx", "fluid": "Oleo"},
    "a002": {"name": "PMGL", "label": "Gas linear", "anp_file": "Gás Linear.xlsx", "fluid": "Gas"},
    "a003": {"name": "PMGD", "label": "Gas diferencial", "anp_file": "Gás Diferencial.xlsx", "fluid": "Gas"},
    "a004": {"name": "PMAE", "label": "Alarmes e eventos", "anp_file": None, "fluid": "Eventos"},
}

OPERATOR_PANEL_EXPORTS = (
    {
        "id": "oil_linear",
        "label": "Óleo Linear",
        "fileName": "Óleo Linear.xlsx",
        "requiredColumns": ["Início Período Medição", "Tag do Ponto Medição", "Volume Bruto Corrigido (m3)"],
        "acceptedColumns": {"Início Período Medição": ["Inicio Período Medição"]},
    },
    {
        "id": "gas_linear",
        "label": "Gás Linear",
        "fileName": "Gás Linear.xlsx",
        "requiredColumns": ["Início Período Medição", "Tag do Ponto Medição", "Volume Bruto Corrigido (m3)"],
        "acceptedColumns": {"Início Período Medição": ["Inicio Período Medição"]},
    },
    {
        "id": "gas_differential",
        "label": "Gás Diferencial",
        "fileName": "Gás Diferencial.xlsx",
        "requiredColumns": ["Início Período Medição", "Tag do Ponto Medição", "Volume Bruto Corrigido (m3)"],
        "acceptedColumns": {"Início Período Medição": ["Inicio Período Medição"]},
    },
    {
        "id": "bsw_inline",
        "label": "BSW em Linha",
        "fileName": "BSW em Linha.xlsx",
        "requiredColumns": ["Data & Hora Medição", "Tag do ponto", "% BSW"],
        "acceptedColumns": {},
    },
    {
        "id": "measurement_failures",
        "label": "Falha de Medição",
        "fileName": "Falha de Medição.xlsx",
        "requiredColumns": ["Código da Falha", "Tipo de Falha", "Tag do Ponto", "Data & Hora da Detecção"],
        "acceptedColumns": {},
    },
)

PARAMETER_RULES = {
    "density_bsw": {
        "label": "Densidade / BSW",
        "parameterKeywords": ["density", "densidade", "bsw", "water cut"],
        "evidenceKeywords": ["density", "densidade", "bsw", "lab", "boletim", "amostra"],
    },
    "pvt": {
        "label": "PVT / propriedades de fluido",
        "parameterKeywords": ["pvt", "viscos", "isentropic", "compress", "z factor", "molecular", "specific heat"],
        "evidenceKeywords": ["pvt", "viscos", "bacpvt", "relatorio pvt"],
    },
    "chromatography": {
        "label": "Cromatografia / composicao do gas",
        "parameterKeywords": [
            "chromat",
            "cromat",
            "composition",
            "composi",
            "methane",
            "ethane",
            "propane",
            "butane",
            "pentane",
            "hexane",
            "heptane",
            "octane",
            "nonane",
            "decane",
            "nitrogen",
            "oxygen",
            "carbon dioxide",
            "carbon monoxide",
            "hydrogen sulphide",
            "hydrogen sulfide",
            "water",
            "etano",
            "metano",
            "heating value",
        ],
        "evidenceKeywords": ["chromat", "cromat", "composition", "composi", "gas analysis", "bacpvtgas", "configuration", "snapshot"],
    },
    "calibration": {
        "label": "Calibracao / fator do medidor",
        "parameterKeywords": ["meter factor", "k factor", "proving", "calibr", "pulse", "meter constant"],
        "evidenceKeywords": ["calib", "calibration", "proving", "corrida", "certificado", "validacao"],
    },
    "pam_limits": {
        "label": "PAM / limites e faixas",
        "parameterKeywords": ["range", "limit", "limite", "alarm limit", "cutoff", "low low", "hi hi", "pressure"],
        "evidenceKeywords": ["pam", "datasheet", "data sheet", "portaria", "memorial", "range", "limite"],
    },
    "uncertainty": {
        "label": "Incerteza de medicao",
        "parameterKeywords": ["uncert", "incerteza"],
        "evidenceKeywords": ["uncert", "incerteza"],
    },
}

EVIDENCE_SOURCE_IDS = (
    "physchem",
    "calibration",
    "uncertainty",
    "equipmentDocs",
    "samplingPlans",
    "regulations",
    "dailyReports",
)


def load_config() -> dict[str, Any]:
    if CONFIG_PATH.exists():
        return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    return {
        "schemaVersion": 1,
        "workspaceRoot": str(ROOT),
        "sources": [
            {
                "id": "workspace",
                "label": "Workspace",
                "kind": "folder",
                "recursive": True,
                "paths": [str(ROOT)],
            }
        ],
    }


CONFIG = load_config()
MATRIX_PATH = APP_DIR / "data" / "matriz_requisitos_sgm1.json"
TEXT_CACHE_PATH = APP_DIR / "data" / "evidence_text_cache.json"


def configured_paths(*source_ids: str) -> list[Path]:
    wanted = set(source_ids)
    paths: list[Path] = []
    for source in CONFIG.get("sources", []):
        if wanted and source.get("id") not in wanted:
            continue
        for raw in source.get("paths", []):
            if not raw:
                continue
            path = Path(raw).expanduser()
            if not path.is_absolute():
                path = Path(CONFIG.get("workspaceRoot") or ROOT) / path
            if path.exists():
                paths.append(path.resolve())
    if not paths:
        paths.append(ROOT.resolve())
    unique: list[Path] = []
    seen = set()
    for path in paths:
        key = str(path).lower()
        if key not in seen:
            seen.add(key)
            unique.append(path)
    return unique


def is_ignored(path: Path) -> bool:
    return ".git" in path.parts or APP_DIR in path.parents


def iter_source_files(source_ids: tuple[str, ...], pattern: str) -> list[Path]:
    files: list[Path] = []
    seen = set()
    for base in configured_paths(*source_ids):
        candidates = [base] if base.is_file() else base.rglob(pattern)
        for item in candidates:
            if item.is_file() and item.match(pattern) and not is_ignored(item):
                key = str(item.resolve()).lower()
                if key not in seen:
                    seen.add(key)
                    files.append(item)
    return sorted(files)


def find_first_file(file_name: str, *source_ids: str) -> Path | None:
    for base in configured_paths(*source_ids):
        if base.is_file() and base.name.lower() == file_name.lower():
            return base
        if base.is_dir():
            direct = base / file_name
            if direct.exists():
                return direct
            for item in base.rglob(file_name):
                if item.is_file() and not is_ignored(item):
                    return item
    return None


def display_path(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT))
    except ValueError:
        return str(path)


def as_number(value: Any) -> float | None:
    if value is None:
        return None
    if isinstance(value, float) and math.isnan(value):
        return None
    text = str(value).strip()
    if not text:
        return None
    text = text.replace(".", "").replace(",", ".") if "," in text else text
    match = re.search(r"[-+]?\d+(?:\.\d+)?", text)
    if not match:
        return None
    try:
        return float(match.group(0))
    except ValueError:
        return None


def fmt_date(value: Any) -> str | None:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return None
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    text = str(value).strip()
    if re.match(r"^\d{4}-\d{2}-\d{2}(?:$|[ T])", text):
        parsed = pd.to_datetime(text, errors="coerce")
    else:
        parsed = pd.to_datetime(text, errors="coerce", dayfirst=True)
    if pd.isna(parsed):
        return None
    return parsed.date().isoformat()


def read_export(path: Path, sheet_name: str | int = 0) -> pd.DataFrame:
    df = pd.read_excel(path, sheet_name=sheet_name)
    if df.empty:
        return df
    first_col = df.columns[0]
    first_values = df[first_col].astype(str).str.lower()
    mask = ~first_values.str.contains("filtros aplicados|nenhum filtro aplicado", na=False)
    df = df[mask].dropna(how="all").copy()
    return df


def first_number(row: pd.Series, candidates: list[str]) -> float | None:
    for col in candidates:
        if col in row.index:
            value = as_number(row.get(col))
            if value is not None:
                return value
    return None


def first_value(row: pd.Series, candidates: list[str]) -> Any:
    for col in candidates:
        if col in row.index and pd.notna(row.get(col)):
            return row.get(col)
    return None


def clean_json(value: Any) -> Any:
    if isinstance(value, dict):
        return {str(key): clean_json(item) for key, item in value.items()}
    if isinstance(value, list):
        return [clean_json(item) for item in value]
    if isinstance(value, tuple):
        return [clean_json(item) for item in value]
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if isinstance(value, float) and math.isnan(value):
        return None
    if pd.isna(value) if not isinstance(value, (dict, list, tuple, str, bytes)) else False:
        return None
    return value


def local_name(tag: str) -> str:
    return tag.split("}", 1)[-1]


def find_text(node: ET.Element, names: list[str]) -> str | None:
    wanted = set(names)
    for child in node.iter():
        if local_name(child.tag) in wanted and child.text and child.text.strip():
            return child.text.strip()
    return None


def op_date_from_path(path: Path) -> str | None:
    for part in path.parts:
        match = re.search(r"Daily reports_(\d{4}-\d{2}-\d{2})", part)
        if match:
            return match.group(1)
    match = re.search(r"_(\d{12,14})_", path.name)
    if match:
        stamp = match.group(1)
        parsed = datetime.strptime(stamp[:8], "%Y%m%d").date()
        return (parsed - timedelta(days=1)).isoformat()
    return None


def parse_xml_bytes(raw: bytes, source: str, op_date: str | None = None) -> dict[str, Any]:
    text = raw.decode("iso-8859-1", errors="replace")
    tree = ET.fromstring(text.encode("iso-8859-1", errors="xmlcharrefreplace"))
    family = local_name(tree.tag).lower()
    records = []
    for node in tree.iter():
        if local_name(node.tag) != "DADOS_BASICOS":
            continue
        tag = node.attrib.get("COD_TAG_PONTO_MEDICAO") or find_text(node, ["COD_TAG_PONTO_MEDICAO"])
        if not tag:
            continue
        volume_corrigido = as_number(
            find_text(node, ["MED_VOLUME_BRTO_CRRGO_MVMDO", "MED_CORRIGIDO_MVMDO"])
        )
        volume_bruto = as_number(find_text(node, ["MED_VOLUME_BRUTO_MVMDO", "MED_BRUTO_MOVIMENTADO"]))
        volume_liquido = as_number(find_text(node, ["MED_VOLUME_LIQUIDO_MVMDO"]))
        records.append(
            {
                "date": op_date,
                "family": family,
                "familyName": FAMILIES.get(family, {}).get("name", family.upper()),
                "tag": tag,
                "volumeCorrigido": volume_corrigido,
                "volumeBruto": volume_bruto,
                "volumeLiquido": volume_liquido,
                "temperatura": as_number(find_text(node, ["MED_TEMPERATURA"])),
                "pressao": as_number(find_text(node, ["MED_PRESSAO_ESTATICA"])),
                "duracaoFluxoMin": as_number(find_text(node, ["PRZ_DURACAO_FLUXO_EFETIVO"])),
                "source": source,
            }
        )
    return {"family": family, "records": records, "recordCount": len(records)}


def parse_measurement_xmls() -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    records: list[dict[str, Any]] = []
    files: list[dict[str, Any]] = []

    for xml_path in iter_source_files(("xmlSent", "dailyReports"), "*.xml"):
        if not re.match(r"00[1-4]_", xml_path.name):
            continue
        op_date = op_date_from_path(xml_path)
        parsed = parse_xml_bytes(xml_path.read_bytes(), display_path(xml_path), op_date)
        if parsed["family"] in FAMILIES:
            records.extend(parsed["records"])
            files.append(
                {
                    "date": op_date,
                    "family": parsed["family"],
                    "familyName": FAMILIES[parsed["family"]]["name"],
                    "path": display_path(xml_path),
                    "kind": "xml",
                    "records": parsed["recordCount"],
                }
            )

    for zip_path in iter_source_files(("xmlSent", "dailyReports"), "*.zip"):
        if not re.match(r"00[1-4]_", zip_path.name):
            continue
        op_date = op_date_from_path(zip_path)
        try:
            zf_handle = zipfile.ZipFile(zip_path)
        except zipfile.BadZipFile:
            print(f"  [AVISO] ZIP corrompido ou inválido, ignorando: {display_path(zip_path)}")
            continue
        with zf_handle as zf:
            for entry in zf.namelist():
                if not entry.lower().endswith(".xml"):
                    continue
                try:
                    xml_bytes = zf.read(entry)
                except Exception:
                    continue
                parsed = parse_xml_bytes(xml_bytes, f"{display_path(zip_path)}::{entry}", op_date)
                family = parsed["family"]
                if family not in FAMILIES:
                    continue
                files.append(
                    {
                        "date": op_date,
                        "family": family,
                        "familyName": FAMILIES[family]["name"],
                        "path": display_path(zip_path),
                        "kind": "zip",
                        "entry": entry,
                        "records": parsed["recordCount"],
                    }
                )
                if family == "a004":
                    # a004 appears only packed in ZIP; keep it as a sent-layer evidence.
                    records.extend(parsed["records"])
    return records, files


def parse_cv_daily() -> dict[tuple[str, str], dict[str, Any]]:
    cv: dict[tuple[str, str], dict[str, Any]] = {}
    for txt_path in iter_source_files(("cvRaw", "dailyReports"), "Run_Daily*.txt"):
        text = txt_path.read_text(encoding="utf-8", errors="ignore")
        meter_match = re.search(r"Meter ID\s+([A-Z0-9\-]+)", text)
        if not meter_match:
            continue
        tag = meter_match.group(1)
        start_match = re.search(r"(?:Start date / time|Period start)\s+(\d{2}/\d{2}/\d{2})", text)
        if not start_match:
            continue
        day = datetime.strptime(start_match.group(1), "%d/%m/%y").date().isoformat()

        def line_value(label: str) -> float | None:
            label_pattern = r"\s+".join(re.escape(part) for part in label.split())
            pattern = rf"^\s*{label_pattern}\s+(?:[A-Za-z0-9/%().-]+\s+)?([-+]?\d+(?:\.\d+)?)"
            match = re.search(pattern, text, re.MULTILINE)
            return float(match.group(1)) if match else None

        gross = line_value("Gross volume")
        base = line_value("Gross standard volume")
        if base is None:
            base = line_value("Base volume")
        net = line_value("Net standard volume")
        item = {
            "date": day,
            "tag": tag,
            "volumeBruto": gross,
            "volumeCorrigido": base,
            "volumeLiquido": net,
            "temperatura": line_value("Meter temperature"),
            "pressao": line_value("Meter pressure"),
            "source": display_path(txt_path),
        }
        current = cv.get((day, tag))
        if current is None or current.get("volumeCorrigido") is None:
            cv[(day, tag)] = item
    return cv


def parse_anp_exports() -> tuple[dict[tuple[str, str, str], dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    index: dict[tuple[str, str, str], dict[str, Any]] = {}
    rows: list[dict[str, Any]] = []
    latest: list[dict[str, Any]] = []

    for family, meta in FAMILIES.items():
        file_name = meta.get("anp_file")
        if not file_name:
            continue
        path = find_first_file(file_name, "anpPanel", "dailyReports")
        if path is None or not path.exists():
            continue
        df = read_export(path)
        start_col = "Início Período Medição" if "Início Período Medição" in df.columns else "Inicio Período Medição"
        vol_corr_col = "Volume Bruto Corrigido (m3)"
        vol_bruto_col = "Volume Bruto (m3)"
        vol_liq_col = "Volume Liquido (m3)"
        temp_col = "Temperatura Fluido  (°C)" if "Temperatura Fluido  (°C)" in df.columns else "Temperatura (°C)"
        press_col = "Pressão Estática (kPa)" if "Pressão Estática (kPa)" in df.columns else "Pressão Estatica (kPa)"
        for _, row in df.iterrows():
            day = fmt_date(row.get(start_col))
            tag = str(row.get("Tag do Ponto Medição", "")).strip()
            if not day or not tag or tag == "nan":
                continue
            item = {
                "date": day,
                "family": family,
                "familyName": meta["name"],
                "tag": tag,
                "volumeCorrigido": as_number(row.get(vol_corr_col)),
                "volumeBruto": as_number(row.get(vol_bruto_col)),
                "volumeLiquido": as_number(row.get(vol_liq_col)),
                "temperatura": as_number(row.get(temp_col)),
                "pressao": as_number(row.get(press_col)),
                "limites": {
                    "temperatura": {
                        "lower": first_number(row, ["Temperatura Inferior Alarme", "Temp Inferior Alarme"]),
                        "upper": first_number(row, ["Temperatura Superior Alarme", "Temp Superior Alarme"]),
                        "fallback": first_number(row, ["Temperatura Adotada Falha", "Temp Adotada Falha"]),
                    },
                    "pressao": {
                        "lower": first_number(row, ["Pressão Limite Inf Alarme", "Limite Inf Alarme", "Pressão Inferior Principal"]),
                        "upper": first_number(row, ["Pressão Limite Sup Alarme", "Limite Sup Alarme", "Pressão Superior Principal"]),
                        "fallback": first_number(row, ["Pressão Adotada Falha", "Pressão Adotada Falha Princ"]),
                    },
                    "diferencial": {
                        "value": first_number(row, ["Diferencial Pressão (kPa)"]),
                        "lower": first_number(row, ["Limite Inferior Alarme", "Pressão Inferior Baixa", "Pressão Inferior Média"]),
                        "upper": first_number(row, ["Limite Superior Alarme", "Pressão Superior Baixa", "Pressão Superior Média"]),
                    },
                    "alarmeHabilitado": first_value(
                        row,
                        [
                            "Alarme Habilitado",
                            "Alarme Pressão Habilitado",
                            "Alarme Temp Habilitado",
                            "Alarme Principal Habilitado",
                        ],
                    ),
                },
                "arquivoCarga": row.get("Nome Arq Carga"),
            }
            index[(day, family, tag)] = item
            rows.append(item)

    latest_by_tag: dict[str, dict[str, Any]] = {}
    for item in rows:
        key = item["tag"]
        if key not in latest_by_tag or item["date"] > latest_by_tag[key]["date"]:
            latest_by_tag[key] = item
    latest = sorted(latest_by_tag.values(), key=lambda x: (x["family"], x["tag"]))
    return index, rows, latest


def _extract_ihm_period_date(rows: list[tuple]) -> str | None:
    """Extract production_date (start of period) from IHM report rows."""
    for row in rows:
        if len(row) > 4 and str(row[2] or "").strip() == "Period:":
            raw = str(row[4] or "")
            # Formato: "06-Jul-2026 00:00 till 07-07-2026 00:00"
            m = re.match(r"(\d{2}-[A-Za-z]{3}-\d{4})", raw)
            if m:
                try:
                    return datetime.strptime(m.group(1), "%d-%b-%Y").date().isoformat()
                except ValueError:
                    pass
            m2 = re.match(r"(\d{2}[-/]\d{2}[-/]\d{4})", raw)
            if m2:
                for fmt in ("%d-%m-%Y", "%d/%m/%Y"):
                    try:
                        return datetime.strptime(m2.group(1), fmt).date().isoformat()
                    except ValueError:
                        pass
    return None


def _extract_ihm_day_totals(rows: list[tuple], fluid: str, source_path: str) -> list[dict[str, Any]]:
    """Extract 'Day totals' records from IHM report rows. Returns list of tag-level records."""
    METRICS = {
        "Gross volume": "gross_volume",
        "Gross standard volume": "gross_standard_volume",
        "Net Standard volume": "net_standard_volume",
        "Mass": "mass",
        "Flow time": "flow_time",
        "Net volume": "net_volume",
        "Energy": "energy",
    }
    UNITS = {
        "gross_volume": "m3",
        "gross_standard_volume": "Sm3",
        "net_standard_volume": "Sm3",
        "mass": "t",
        "flow_time": "min",
        "net_volume": "m3",
        "energy": "GJ",
    }

    results: list[dict[str, Any]] = []
    i = 0
    current_skid: str = ""
    while i < len(rows):
        row = rows[i]
        vals = [str(c or "").strip() for c in row]

        # Detectar nome do skid (linha com texto em col[5..] que não é métrica)
        if vals[2] and not any(vals[2].startswith(m) for m in ("Day totals", "Cumulative", "Flow weighted", "Period")):
            skid_candidate = " | ".join(c for c in vals[5:] if c and c not in ("Run 1", "Run 2", "Run 3", "Total", "-"))
            if skid_candidate:
                current_skid = skid_candidate

        # Encontrar seção "Day totals"
        if vals[2] == "Day totals" or (vals[2] == "" and vals[3] == "" and any(v == "Day totals" for v in vals)):
            i += 1
            if i >= len(rows):
                break
            tag_row = rows[i]
            # Construir mapeamento col_idx -> tag, ignorando None/'Run N'/'Total'
            tag_cols: dict[int, str] = {}
            for ci, cell in enumerate(tag_row):
                c = str(cell or "").strip()
                if c and c not in ("Run 1", "Run 2", "Run 3", "Run 4", "Total", "-", "None"):
                    tag_cols[ci] = c

            if not tag_cols:
                i += 1
                continue

            # Ler métricas até encontrar linha vazia ou nova seção
            metrics_data: dict[str, dict[int, float]] = {}
            i += 1
            while i < len(rows):
                mrow = rows[i]
                mvals = [str(c or "").strip() for c in mrow]
                metric_name = mvals[3] if len(mvals) > 3 else ""
                if not metric_name:
                    # Linha vazia — fim desta seção
                    if all(not v for v in mvals):
                        break
                    i += 1
                    continue
                if metric_name in ("Day totals", "Cumulative totals @ day close", "Flow weighted averages"):
                    break
                if metric_name in METRICS:
                    field = METRICS[metric_name]
                    metrics_data[field] = {}
                    for ci, tag in tag_cols.items():
                        if ci < len(mrow):
                            try:
                                v = float(mrow[ci]) if mrow[ci] not in (None, "", "-") else 0.0
                                metrics_data[field][ci] = v
                            except (ValueError, TypeError):
                                metrics_data[field][ci] = 0.0
                i += 1

            # Montar registros por tag
            for ci, tag in tag_cols.items():
                record: dict[str, Any] = {
                    "fluid": fluid,
                    "tag": tag,
                    "skid": current_skid,
                    "source": str(Path(source_path).name),
                }
                for field, col_vals in metrics_data.items():
                    record[field] = col_vals.get(ci, 0.0)
                    record[f"{field}_unit"] = UNITS.get(field, "")
                if any(record.get(f, 0) for f in METRICS.values() if f in record):
                    results.append(record)
            continue
        i += 1
    return results


def parse_ihm_daily_reports() -> dict[str, Any]:
    """Parse Daily_Oil, Daily_Gas, Daily_Water IHM reports from all daily packages."""
    try:
        from openpyxl import load_workbook  # noqa: PLC0415
    except ImportError:
        return {"rows": [], "summary": {}, "error": "openpyxl não instalado"}

    fluid_patterns = [
        ("oil",   "Daily_Oil*.xlsx"),
        ("gas",   "Daily_Gas*.xlsx"),
        ("water", "Daily_Water*.xlsx"),
    ]

    all_rows: list[dict[str, Any]] = []
    errors: list[str] = []

    for fluid, pattern in fluid_patterns:
        for xlsx_path in iter_source_files(("dailyReports",), pattern):
            try:
                wb = load_workbook(str(xlsx_path), read_only=True, data_only=True)
            except Exception as exc:
                errors.append(f"{xlsx_path.name}: {exc}")
                continue
            try:
                ws = wb.active
                rows = list(ws.iter_rows(values_only=True))
            except Exception as exc:
                errors.append(f"{xlsx_path.name}: {exc}")
                continue
            finally:
                wb.close()

            production_date = _extract_ihm_period_date(rows)
            if not production_date:
                # Tenta extrair do nome do arquivo: Daily_Oil_2026-07-07.xlsx → 2026-07-06
                m = re.search(r"(\d{4}-\d{2}-\d{2})", xlsx_path.name)
                if m:
                    try:
                        d = date.fromisoformat(m.group(1))
                        production_date = (d - timedelta(days=1)).isoformat()
                    except ValueError:
                        pass

            if not production_date:
                errors.append(f"{xlsx_path.name}: data de produção não identificada")
                continue

            records = _extract_ihm_day_totals(rows, fluid, str(xlsx_path))
            for rec in records:
                rec["production_date"] = production_date
            all_rows.extend(records)

    # Sumário
    days = sorted({r["production_date"] for r in all_rows})
    by_fluid: dict[str, int] = {}
    by_day: dict[str, dict[str, float]] = {}
    for r in all_rows:
        by_fluid[r["fluid"]] = by_fluid.get(r["fluid"], 0) + 1
        day = r["production_date"]
        if day not in by_day:
            by_day[day] = {"oil_gsv_sm3": 0.0, "gas_gv_m3": 0.0, "water_gv_m3": 0.0}
        fluid = r["fluid"]
        gv = r.get("gross_standard_volume") or r.get("gross_volume") or 0.0
        if fluid == "oil":
            by_day[day]["oil_gsv_sm3"] += gv
        elif fluid == "gas":
            by_day[day]["gas_gv_m3"] += gv
        elif fluid == "water":
            by_day[day]["water_gv_m3"] += gv

    return {
        "rows": all_rows,
        "days": [
            {"date": d, **by_day.get(d, {})}
            for d in days
        ],
        "summary": {
            "total_records": len(all_rows),
            "days": len(days),
            "by_fluid": by_fluid,
            "errors": len(errors),
        },
        "errors": errors,
    }


def parse_gas_balance() -> dict[str, Any]:
    """Parse GasBalance IHM report from all daily packages."""
    try:
        from openpyxl import load_workbook  # noqa: PLC0415
    except ImportError:
        return {"rows": [], "summary": {}}

    BALANCE_LABELS = {
        "Fiscal gas metering": "fiscal_gas_m3",
        "Gas lift injection": "gas_lift_m3",
        "HP flare": "hp_flare_m3",
        "LP flare": "lp_flare_m3",
        "Fuel gas": "fuel_gas_m3",
        "Gas injection": "gas_injection_m3",
        "Pilot": "pilot_gas_m3",
        "Vent": "vent_gas_m3",
        "Net production": "net_production_m3",
    }

    all_rows: list[dict[str, Any]] = []

    for xlsx_path in iter_source_files(("dailyReports",), "GasBalance*.xlsx"):
        try:
            wb = load_workbook(str(xlsx_path), read_only=True, data_only=True)
            ws = wb.active
            rows = list(ws.iter_rows(values_only=True))
            wb.close()
        except Exception:
            continue

        production_date = _extract_ihm_period_date(rows)
        if not production_date:
            m = re.search(r"(\d{4}-\d{2}-\d{2})", xlsx_path.name)
            if m:
                try:
                    d = date.fromisoformat(m.group(1))
                    production_date = (d - timedelta(days=1)).isoformat()
                except ValueError:
                    pass
        if not production_date:
            continue

        record: dict[str, Any] = {"production_date": production_date, "source": xlsx_path.name}
        for row in rows:
            label = str(row[2] or row[3] or "").strip()
            for key, field in BALANCE_LABELS.items():
                if key.lower() in label.lower():
                    # Procurar primeiro valor numérico na linha
                    for cell in row[4:]:
                        try:
                            record[field] = float(cell)
                            break
                        except (TypeError, ValueError):
                            pass
        if len(record) > 2:
            all_rows.append(record)

    days = sorted({r["production_date"] for r in all_rows})
    return {
        "rows": all_rows,
        "summary": {"days": len(days), "total_records": len(all_rows)},
    }


def parse_points() -> dict[str, dict[str, Any]]:
    path = find_first_file("Pontos de Medição.xlsx", "cadastro", "anpPanel")
    if path is None or not path.exists():
        return {}
    df = read_export(path)
    points = {}
    for _, row in df.iterrows():
        tag = str(row.get("TAG_PONTO_MEDICAO", "")).strip()
        if not tag or tag == "nan":
            continue
        points[tag] = {
            "tag": tag,
            "fluid": row.get("FLUIDO"),
            "principal": row.get("TIPO_MEDICAO_PRINCIPAL"),
            "secondary": row.get("TIPO_MEDICAO_SECUNDARIA"),
            "meterType": row.get("TIPO_MEDIDOR"),
            "active": row.get("IND_ATIVO"),
            "minOperacao": as_number(row.get("MINIMO_OPERACAO")),
            "maxOperacao": as_number(row.get("MAXIMO_OPERACAO")),
            "incertezaMaxima": as_number(row.get("INCERTEZA_MAXIMA")),
            "computadorVazao": row.get("COMPUTADOR_VAZAO"),
        }
    return points


def parse_bsw() -> tuple[list[dict[str, Any]], dict[str, Any]]:
    path = find_first_file("BSW em Linha.xlsx", "anpPanel", "dailyReports")
    if path is None or not path.exists():
        return [], {}
    df = read_export(path)
    rows = []
    for _, row in df.iterrows():
        rows.append(
            {
                "date": fmt_date(row.get("Data & Hora Medição")),
                "tag": row.get("Tag do ponto"),
                "bsw": as_number(row.get("% BSW")),
                "maxBsw": as_number(row.get("% Máximo BSW")),
                "arquivo": row.get("Arquivo Recebido"),
            }
        )
    max_row = max(rows, key=lambda x: (x.get("bsw") or 0), default={})
    return rows, max_row


# ── Mapeamento de categorias de eventos regulatórios a partir do nome da pasta ──
_CARTA_CATEGORY_MAP = [
    ("aprovac",       "system_approval",     "Aprovação de Sistema"),
    ("est-vent",      "est_vent_tank",        "EST-VENT-TANK"),
    ("auditoria",     "audit",               "Auditoria ANP"),
    ("calibra",       "calibration",         "Calibração"),
    ("falha acima",   "failure_240h",        "Falha >240h"),
    ("falha de",      "failure_240h",        "Falha >240h"),
    ("nfsm",          "nfsm",                "NFSM"),
    ("240",           "failure_240h",        "Falha >240h"),
    ("mpfm",          "mpfm_approval",       "Aprovação MPFM"),
    ("multifasic",    "mpfm_approval",       "Aprovação MPFM"),
    ("operac",        "mpfm_authorization",  "Autorização Operação"),
    ("corrente",      "oil_spec",            "Corrente de Petróleo"),
    ("pev",           "oil_spec",            "Corrente de Petróleo"),
    ("flowline",      "flowline",            "Flowline"),
    ("offloading",    "offloading",          "Offloading"),
    ("nc ",           "nc_closure",          "Encerramento NC"),
    ("encerramento",  "nc_closure",          "Encerramento NC"),
    ("dp ",           "cv_parameter",        "Parâmetro CV"),
    ("escala",        "cv_parameter",        "Parâmetro CV"),
]

_TAG_PATTERN = re.compile(r"\b(\d{2}[A-Z]{2}\d{4}[A-Z]?)\b")
_DOC_PATTERN = re.compile(r"Equinor-BRA-(\d{4})-?(\d+)", re.IGNORECASE)
_DATE_PATTERN = re.compile(r"(20\d{2})[_\-](\d{2})[_\-](\d{2})")
_OFICIO_PATTERN = re.compile(r"Oficio[_\- ](\d+)-(\d{4})", re.IGNORECASE)


def _infer_carta_category(text: str) -> tuple[str, str]:
    lower = text.lower()
    for kw, code, label in _CARTA_CATEGORY_MAP:
        if kw in lower:
            return code, label
    return "regulatory", "Regulatório"


def _extract_carta_meta(folder_name: str, files: list[str]) -> dict[str, Any] | None:
    all_text = folder_name + " " + " ".join(files)
    # Data
    dm = _DATE_PATTERN.search(all_text)
    event_date = f"{dm.group(1)}-{dm.group(2)}-{dm.group(3)}" if dm else None
    # Código Equinor
    doc_m = _DOC_PATTERN.search(all_text)
    doc_number = f"Equinor-BRA-{doc_m.group(1)}-{doc_m.group(2)}" if doc_m else None
    # Ofício ANP
    oficio_m = _OFICIO_PATTERN.search(all_text)
    oficio = f"Ofício {oficio_m.group(1)}/{oficio_m.group(2)}" if oficio_m else None
    # Tags
    tags = list({m.group(1) for m in _TAG_PATTERN.finditer(all_text)})
    # Categoria
    category_code, category_label = _infer_carta_category(folder_name)
    # Status heurístico
    lower = folder_name.lower() + " ".join(files).lower()
    status = "closed" if any(w in lower for w in ["encerramento", "aprovac", "despacho", "oficio 7", "oficio 9"]) else "open"
    return {
        "folder": folder_name,
        "event_date": event_date,
        "category": category_code,
        "category_label": category_label,
        "doc_number": doc_number,
        "oficio": oficio,
        "tags": sorted(tags),
        "status": status,
        "title": folder_name.lstrip("0123456789 -"),
    }


def parse_cartas_anp() -> dict[str, Any]:
    """Extrai metadados de eventos regulatórios das pastas de Cartas e Ofícios ANP."""
    events: list[dict[str, Any]] = []
    for carta_path in configured_paths("cartasAnp"):
        if not carta_path.is_dir():
            continue
        # Pasta raiz: cada subdiretório = um evento
        for folder in sorted(carta_path.iterdir()):
            if not folder.is_dir() or folder.name.startswith("00 - Modelos"):
                continue
            files = [f.name for f in folder.iterdir() if f.is_file()] if folder.exists() else []
            meta = _extract_carta_meta(folder.name, files)
            if meta:
                events.append(meta)

    events.sort(key=lambda e: (e.get("event_date") or ""), reverse=True)
    open_events = [e for e in events if e.get("status") == "open"]
    return {
        "events": events,
        "summary": {
            "total": len(events),
            "open": len(open_events),
            "by_category": {
                k: sum(1 for e in events if e["category"] == k)
                for k in {e["category"] for e in events}
            },
        },
    }


def parse_open_nfsms() -> list[dict[str, Any]]:
    """
    Retorna lista de NFSMs abertas conhecidas a partir das pastas de Cartas e Ofícios
    e de informações fixas do handover. Complementa os dados da planilha Falha de Medição.
    """
    known_open: list[dict[str, Any]] = [
        {
            "tag": "43FT0102",
            "name": "HP Flare Gas",
            "issue": "Picos intermitentes de alta vazão",
            "detection_date": "2026-03-01",
            "nfsm_ref": "BACALHAU 2026.033",
            "status": "open",
            "notes": "Falha T enviada mensalmente. 2º embarque FB Solutions aguardado.",
        },
        {
            "tag": "45FT0555",
            "name": "Fuel Gas Total (CTG)",
            "issue": "Erro na escala do DP (escala incorreta configurada no CV)",
            "detection_date": "2026-05-07",
            "nfsm_ref": "Carta 18 - Equinor-BRA-2026-xxx",
            "status": "open",
            "notes": "Notificação de falha enviada em 07/05/2026. Verificar correção dos volumes afetados.",
        },
    ]
    return known_open


def parse_alarm_management() -> dict[str, Any]:
    """
    Lê os arquivos de Gestão de Alarmes (xlsm) da aba AlarmesConsolidado.
    Extrai alarmes de medição relevantes com data, tag, tipo e status.
    """
    try:
        from openpyxl import load_workbook  # noqa: PLC0415
    except ImportError:
        return {"rows": [], "summary": {}}

    METERING_KEYWORDS = [
        "falha de pulso", "pulse fail", "failure", "falha", "medicao", "medição",
        "comunicacao", "communication", "sensor", "transmitter", "transm",
        "bsw", "densidad", "pressão", "pressure", "temperatura", "temperature",
        "vazão", "flow", "medidor", "meter", "calibra",
    ]

    all_rows: list[dict[str, Any]] = []
    for src in configured_paths("gestaoAlarmes"):
        for xlsm_path in sorted(src.glob("*.xlsm"), key=lambda f: f.stat().st_size, reverse=True):
            try:
                wb = load_workbook(str(xlsm_path), read_only=True, data_only=True)
            except Exception:
                continue
            for sheet_name in ("AlarmesConsolidado", "PAINEL", wb.active.title):
                if sheet_name not in wb.sheetnames:
                    continue
                ws = wb[sheet_name]
                raw = list(ws.iter_rows(values_only=True))
                # Procurar header row
                h_idx = next(
                    (i for i, r in enumerate(raw[:10])
                     if any(str(c or "").strip() != "" for c in r[:5])),
                    0,
                )
                headers = [str(c or "").strip().lower() for c in raw[h_idx]]

                def _col(kw: str) -> int | None:
                    for i, h in enumerate(headers):
                        if kw in h:
                            return i
                    return None

                date_col = _col("data") or _col("date") or 0
                tag_col  = _col("tag") or _col("cv") or 1
                alarm_col = _col("alarm") or _col("alarme") or _col("descrição") or 2
                sev_col  = _col("severidade") or _col("severity") or _col("nivel") or 3
                stat_col = _col("status") or _col("estado") or 4

                for row in raw[h_idx + 1:]:
                    if not any(row[:5]):
                        continue
                    desc = str(row[alarm_col] if alarm_col < len(row) else "").lower()
                    # Filtrar apenas alarmes relacionados a medição
                    if not any(kw in desc for kw in METERING_KEYWORDS):
                        continue
                    def _fmt(v: Any) -> str:
                        if v is None:
                            return ""
                        if hasattr(v, "isoformat"):
                            return v.date().isoformat() if hasattr(v, "date") else v.isoformat()[:10]
                        return str(v).strip()
                    all_rows.append({
                        "date":     _fmt(row[date_col] if date_col < len(row) else None),
                        "tag":      _fmt(row[tag_col]  if tag_col  < len(row) else None),
                        "alarm":    _fmt(row[alarm_col] if alarm_col < len(row) else None),
                        "severity": _fmt(row[sev_col]  if sev_col  < len(row) else None),
                        "status":   _fmt(row[stat_col] if stat_col  < len(row) else None),
                        "source":   xlsm_path.name,
                    })
                break  # primeira aba válida
            wb.close()

    all_rows.sort(key=lambda r: r.get("date") or "", reverse=True)
    by_tag: dict[str, int] = {}
    for r in all_rows:
        by_tag[r["tag"]] = by_tag.get(r["tag"], 0) + 1

    return {
        "rows": all_rows[:200],
        "summary": {
            "total": len(all_rows),
            "by_tag": by_tag,
        },
    }


def parse_sfp_registration() -> dict[str, Any]:
    """
    Lê o arquivo mais completo de Cadastro no SFP (versão _Mariana ou maior numeração).
    Extrai: tag do ponto, em operação, número de série, possui retificador, etc.
    """
    try:
        from openpyxl import load_workbook  # noqa: PLC0415
    except ImportError:
        return {"rows": [], "summary": {}}

    # Pegar o arquivo com mais linhas na pasta
    best_path: "Path | None" = None
    best_count = 0
    for src in configured_paths("sfpRegistration"):
        for f in sorted(src.glob("*.xlsx")):
            try:
                wb = load_workbook(str(f), read_only=True, data_only=True)
                ws = wb.active
                count = sum(1 for r in ws.iter_rows(values_only=True) if any(c for c in list(r)[:5]))
                wb.close()
                if count > best_count:
                    best_count = count
                    best_path = f
            except Exception:
                pass

    if best_path is None:
        return {"rows": [], "summary": {"error": "Arquivo não encontrado"}}

    try:
        wb = load_workbook(str(best_path), read_only=True, data_only=True)
        ws = wb.active
        raw_rows = list(ws.iter_rows(values_only=True))
        wb.close()
    except Exception as exc:
        return {"rows": [], "summary": {"error": str(exc)}}

    rows: list[dict[str, Any]] = []
    current_tag = ""
    for row in raw_rows:
        if len(row) < 5:
            continue
        tag = str(row[1] or "").strip()
        if tag and tag not in ("Tag do Ponto de Medição", "Identificação no sistema"):
            current_tag = tag
            meter_tag = str(row[2] or "").strip()
            in_service = str(row[3] or "").strip().lower()
            serial = str(row[4] or "").strip()
            rows.append({
                "tag": current_tag,
                "meter_tag": meter_tag,
                "in_service": in_service in ("sim", "yes", "s", "y", "1"),
                "serial_number": serial,
                "flow_conditioner": str(row[7] or "").strip() if len(row) > 7 else "",
                "upstream_tag": str(row[8] or "").strip() if len(row) > 8 else "",
            })
        elif not tag and current_tag and len(row) > 2:
            # Sub-linha do mesmo ponto (outros transmissores)
            meter_tag = str(row[2] or "").strip()
            in_service = str(row[3] or "").strip().lower()
            serial = str(row[4] or "").strip()
            if meter_tag and meter_tag not in ("Tag do Medidor",):
                rows.append({
                    "tag": current_tag,
                    "meter_tag": meter_tag,
                    "in_service": in_service in ("sim", "yes", "s", "y", "1"),
                    "serial_number": serial,
                    "flow_conditioner": str(row[7] or "").strip() if len(row) > 7 else "",
                    "upstream_tag": str(row[8] or "").strip() if len(row) > 8 else "",
                })

    unique_tags = sorted({r["tag"] for r in rows})
    in_service = [r for r in rows if r["in_service"]]
    return {
        "rows": rows,
        "source": str(best_path),
        "summary": {
            "unique_tags": len(unique_tags),
            "total_elements": len(rows),
            "in_service": len(in_service),
            "tag_list": unique_tags,
        },
    }


def parse_operating_ranges() -> dict[str, Any]:
    """
    Lê Process condition FPSO Bacalhau.xlsx e retorna faixas operacionais por TAG.
    Colunas: TAG, Sistema, Min/Ops/Max Flow, Pres, Temp, Density, Molecular Weight
    """
    try:
        from openpyxl import load_workbook  # noqa: PLC0415
    except ImportError:
        return {"rows": [], "summary": {}}

    path = find_first_file("Process condition FPSO Bacalhau.xlsx", "operatingRanges", "equipmentDocs")
    if path is None or not path.exists():
        for src in configured_paths("operatingRanges"):
            for f in src.rglob("Process condition*.xlsx"):
                path = f
                break
    if path is None or not path.exists():
        return {"rows": [], "summary": {"error": "Arquivo não encontrado"}}

    try:
        wb = load_workbook(str(path), read_only=True, data_only=True)
        ws = wb.active
        raw_rows = list(ws.iter_rows(values_only=True))
        wb.close()
    except Exception as exc:
        return {"rows": [], "summary": {"error": str(exc)}}

    # Encontrar header row
    header_idx = next(
        (i for i, r in enumerate(raw_rows)
         if any(str(c or "").strip().upper() in ("TAG", "ITEM") for c in r[:5])),
        None,
    )
    if header_idx is None:
        return {"rows": [], "summary": {"error": "Header não encontrado"}}

    headers = [str(c or "").strip() for c in raw_rows[header_idx]]

    def _num(v: Any) -> float | None:
        if v is None:
            return None
        try:
            return float(str(v).replace(",", ".").strip())
        except (ValueError, TypeError):
            return None

    rows: list[dict[str, Any]] = []
    for row in raw_rows[header_idx + 1:]:
        tag = str(row[3] or "").strip() if len(row) > 3 else ""
        if not tag:
            continue
        rows.append({
            "tag":        tag,
            "application": str(row[1] or "").strip(),
            "service":    str(row[2] or "").strip(),
            "system":     str(row[4] or "").strip() if len(row) > 4 else "",
            "flow_min":   _num(row[5]  if len(row) > 5 else None),
            "flow_ops":   _num(row[6]  if len(row) > 6 else None),
            "flow_max":   _num(row[7]  if len(row) > 7 else None),
            "flow_unit":  str(row[8] or "").strip() if len(row) > 8 else "",
            "pres_min":   _num(row[9]  if len(row) > 9 else None),
            "pres_ops":   _num(row[10] if len(row) > 10 else None),
            "pres_max":   _num(row[11] if len(row) > 11 else None),
            "pres_unit":  str(row[12] or "").strip() if len(row) > 12 else "",
            "temp_min":   _num(row[13] if len(row) > 13 else None),
            "temp_ops":   _num(row[14] if len(row) > 14 else None),
            "temp_max":   _num(row[15] if len(row) > 15 else None),
            "temp_unit":  str(row[16] or "").strip() if len(row) > 16 else "",
            "density":    _num(row[17] if len(row) > 17 else None),
            "dens_unit":  str(row[18] or "").strip() if len(row) > 18 else "",
            "mol_weight": _num(row[19] if len(row) > 19 else None),
        })

    # Indexar por TAG para acesso rápido
    by_tag = {r["tag"]: r for r in rows}
    return {
        "rows": rows,
        "by_tag": by_tag,
        "source": str(path),
        "summary": {"total": len(rows), "tags": [r["tag"] for r in rows]},
    }


def parse_calibration_control() -> dict[str, Any]:
    """
    Lê Calibration Control - Primários.xlsx e retorna status de calibração por TAG.
    Colunas: Classe, Sistema, TAG, Série, Equipamento, Faixa, Certificado,
             Próxima Calibração, Periodicidade(dias), Status, Dias p/Vencimento, Última Calibração
    """
    try:
        from openpyxl import load_workbook  # noqa: PLC0415
    except ImportError:
        return {"rows": [], "summary": {}}

    rows: list[dict[str, Any]] = []
    path = find_first_file("Calibration Control - Primários.xlsx", "calibrationControl")
    if path is None or not path.exists():
        # Fallback: busca qualquer Calibration Control em subpastas
        for src_path in configured_paths("calibrationControl"):
            for f in src_path.rglob("Calibration Control*Prim*.xlsx"):
                path = f
                break
    if path is None or not path.exists():
        return {"rows": [], "summary": {"error": "Arquivo não encontrado", "path": None}}

    try:
        wb = load_workbook(str(path), read_only=True, data_only=True)
        ws = wb.active
        raw_rows = list(ws.iter_rows(values_only=True))
        wb.close()
    except Exception as exc:
        return {"rows": [], "summary": {"error": str(exc)}}

    # Encontrar linha de headers (contém "TAG")
    header_row_idx = next((i for i, r in enumerate(raw_rows) if any("TAG" == str(c or "").strip().upper() for c in r)), None)
    if header_row_idx is None:
        return {"rows": [], "summary": {"error": "Header não encontrado"}}

    headers = [str(c or "").strip() for c in raw_rows[header_row_idx]]
    COL = {h: i for i, h in enumerate(headers)}
    # Nomes esperados (normalizar)
    def col(name: str) -> int | None:
        for k, v in COL.items():
            if name.lower() in k.lower():
                return v
        return None

    tag_col      = col("tag")
    system_col   = col("sistema")
    class_col    = col("classe")
    equip_col    = col("equipamento")
    cert_col     = col("certificado")
    next_cal_col = col("próxima")
    period_col   = col("periodicidade")
    status_col   = col("status")
    days_col     = col("dias")
    last_cal_col = col("última")

    if tag_col is None:
        return {"rows": [], "summary": {"error": "Coluna TAG não encontrada"}}

    last_class = last_system = last_tag = ""
    for row in raw_rows[header_row_idx + 1:]:
        tag = str(row[tag_col] or "").strip()
        if not tag and last_tag:
            tag = last_tag  # mesma tag, linha de faixa adicional → skip
            continue
        klass = str(row[class_col] or "").strip() if class_col is not None else last_class
        system = str(row[system_col] or "").strip() if system_col is not None else last_system
        if klass:
            last_class = klass
        else:
            klass = last_class
        if system:
            last_system = system
        else:
            system = last_system
        if not tag:
            continue
        last_tag = tag

        next_cal = row[next_cal_col] if next_cal_col is not None else None
        last_cal = row[last_cal_col] if last_cal_col is not None else None
        status = str(row[status_col] or "").strip() if status_col is not None else ""
        days_to_exp = None
        if days_col is not None:
            try:
                days_to_exp = int(row[days_col]) if row[days_col] is not None else None
            except (ValueError, TypeError):
                pass
        period = None
        if period_col is not None:
            try:
                period = int(row[period_col]) if row[period_col] is not None else None
            except (ValueError, TypeError):
                pass

        def _fmt_dt(v: Any) -> str | None:
            if v is None:
                return None
            if hasattr(v, "isoformat"):
                return v.date().isoformat() if hasattr(v, "date") else v.isoformat()[:10]
            return str(v)[:10]

        rows.append({
            "tag": tag,
            "system": system or last_system,
            "classification": klass or last_class,
            "equipment": str(row[equip_col] or "").strip() if equip_col is not None else "",
            "certificate": str(row[cert_col] or "").strip() if cert_col is not None else "",
            "last_calibration": _fmt_dt(last_cal),
            "next_calibration": _fmt_dt(next_cal),
            "period_days": period,
            "status": status,
            "days_to_expiry": days_to_exp,
            "overdue": (days_to_exp is not None and days_to_exp < 0),
            "expiry_soon": (days_to_exp is not None and 0 <= days_to_exp <= 30),
        })

    overdue = [r for r in rows if r["overdue"]]
    soon    = [r for r in rows if r["expiry_soon"]]
    return {
        "rows": rows,
        "source": str(path),
        "summary": {
            "total": len(rows),
            "overdue": len(overdue),
            "expiry_soon": len(soon),
            "ok": len([r for r in rows if not r["overdue"] and not r["expiry_soon"]]),
            "overdue_tags": [r["tag"] for r in overdue],
        },
    }


def parse_failures() -> dict[str, Any]:
    path = find_first_file("Falha de Medição.xlsx", "anpPanel", "dailyReports")
    if path is None or not path.exists():
        return {"total": 0, "open": 0, "byType": [], "latestOpen": [], "knownOpen": parse_open_nfsms()}
    df = read_export(path)
    returned = df["Data & Hora do Retorno"].notna() if "Data & Hora do Retorno" in df.columns else pd.Series([False] * len(df))
    df_open = df[~returned].copy()
    type_counts = Counter(str(v) for v in df.get("Tipo de Falha", pd.Series(dtype=str)).dropna())
    latest = []
    for _, row in df_open.iterrows():
        detected = fmt_date(row.get("Data & Hora da Detecção"))
        due_days = as_number(row.get("Previsão de Retorna em Dias")) or 0
        due_date = None
        overdue_days = None
        if detected:
            due = datetime.fromisoformat(detected).date() + timedelta(days=int(due_days))
            due_date = due.isoformat()
            overdue_days = (TODAY - due).days
        latest.append(
            {
                "code": row.get("Código da Falha"),
                "type": row.get("Tipo de Falha"),
                "notification": row.get("Tipo de Notificação"),
                "tag": row.get("Tag do Ponto"),
                "detected": detected,
                "dueDate": due_date,
                "overdueDays": overdue_days,
                "arquivo": row.get("Arquivo Recebido"),
            }
        )
    latest.sort(key=lambda x: (x.get("overdueDays") is None, -(x.get("overdueDays") or -99999)))

    parecer = {}
    parecer_path = find_first_file("Parecer.xlsx", "anpPanel", "dailyReports")
    if parecer_path is not None and parecer_path.exists():
        p_df = read_export(parecer_path)
        col = "Parecer da ANP"
        if col in p_df.columns:
            parecer = dict(Counter(str(v) for v in p_df[col].dropna()))

    return {
        "total": int(len(df)),
        "open": int(len(df_open)),
        "returned": int(returned.sum()),
        "byType": [{"name": key, "value": value} for key, value in type_counts.most_common()],
        "parecer": [{"name": key, "value": value} for key, value in parecer.items()],
        "latestOpen": latest[:10],
    }


def column_exists(columns: list[str], expected: str, accepted: dict[str, list[str]]) -> bool:
    candidates = [expected, *(accepted.get(expected) or [])]
    normalized = {str(column).strip().lower() for column in columns}
    return any(candidate.strip().lower() in normalized for candidate in candidates)


def inspect_operator_panel_export(export: dict[str, Any]) -> dict[str, Any]:
    path = find_first_file(export["fileName"], "anpPanel", "dailyReports")
    if path is None or not path.exists():
        return {
            "id": export["id"],
            "label": export["label"],
            "fileName": export["fileName"],
            "status": "critical",
            "path": None,
            "rows": 0,
            "missingColumns": export["requiredColumns"],
            "latestDate": None,
            "message": "Arquivo obrigatório não localizado no Painel do Operador.",
        }
    try:
        df = read_export(path)
    except Exception as exc:
        return {
            "id": export["id"],
            "label": export["label"],
            "fileName": export["fileName"],
            "status": "critical",
            "path": display_path(path),
            "rows": 0,
            "missingColumns": [],
            "latestDate": None,
            "message": f"Falha ao ler arquivo: {exc.__class__.__name__}.",
        }
    columns = [str(column).strip() for column in df.columns]
    accepted = export.get("acceptedColumns") or {}
    missing_columns = [column for column in export["requiredColumns"] if not column_exists(columns, column, accepted)]
    date_columns = ["Início Período Medição", "Inicio Período Medição", "Data & Hora Medição", "Data & Hora da Detecção"]
    latest_date = None
    for date_column in date_columns:
        if date_column not in df.columns:
            continue
        values = [fmt_date(value) for value in df[date_column].dropna().tail(500)]
        latest_date = max((value for value in values if value), default=None)
        if latest_date:
            break
    status = "critical" if missing_columns or df.empty else "ok"
    message = "Arquivo pronto para ingestão prioritária."
    if missing_columns:
        message = "Colunas obrigatórias ausentes: " + ", ".join(missing_columns)
    elif df.empty:
        message = "Arquivo localizado, mas sem linhas úteis após limpeza."
    return {
        "id": export["id"],
        "label": export["label"],
        "fileName": export["fileName"],
        "status": status,
        "path": display_path(path),
        "rows": int(len(df)),
        "missingColumns": missing_columns,
        "latestDate": latest_date,
        "message": message,
    }


def build_operator_panel_health() -> dict[str, Any]:
    exports = [inspect_operator_panel_export(export) for export in OPERATOR_PANEL_EXPORTS]
    missing_files = [item for item in exports if not item.get("path")]
    missing_info = [item for item in exports if item.get("path") and item.get("status") != "ok"]
    status = "critical" if missing_files or missing_info else "ok"
    return {
        "priority": "alta",
        "sourceId": "anpPanel",
        "label": "Painel do Operador",
        "status": status,
        "required": len(exports),
        "ready": sum(1 for item in exports if item.get("status") == "ok"),
        "missingFiles": len(missing_files),
        "missingInformation": len(missing_info),
        "exports": exports,
        "message": "Prioridade de ingestão atendida." if status == "ok" else "Há arquivo ou informação obrigatória do Painel do Operador pendente.",
    }


def parse_mpfm() -> dict[str, Any]:
    # Tenta encontrar o arquivo mensal MPFM do mês atual ou mais recente disponível
    MONTH_NAMES_PT = {
        1: "JAN", 2: "FEV", 3: "MAR", 4: "ABR", 5: "MAI", 6: "JUN",
        7: "JUL", 8: "AGO", 9: "SET", 10: "OUT", 11: "NOV", 12: "DEZ",
    }
    MONTH_NAMES_EN = {
        1: "JAN", 2: "FEB", 3: "MAR", 4: "APR", 5: "MAY", 6: "JUN",
        7: "JUL", 8: "AUG", 9: "SEP", 10: "OCT", 11: "NOV", 12: "DEC",
    }
    candidates = []
    for m in range(TODAY.month, 0, -1):
        y = TODAY.year
        for mon in (MONTH_NAMES_PT[m], MONTH_NAMES_EN[m], f"{m:02d}"):
            candidates += [
                f"MPFM_{mon}_{y}.xlsx",
                f"MPFM_{mon}_{str(y)[2:]}.xlsx",
                f"MPFM_{mon.capitalize()}_{y}.xlsx",
            ]
    # Fallback legado
    candidates += ["MPFM_JUN_2026.xlsx", "MPFM_JUN_26.xlsx"]

    path = None
    for name in candidates:
        found = find_first_file(name, "mpfm", "dailyReports")
        if found and found.exists():
            path = found
            break

    if path is None or not path.exists():
        return {"status": [], "alerts": [], "source": None}

    try:
        status_df = pd.read_excel(path, sheet_name="STATUS_MES")
        alert_df = pd.read_excel(path, sheet_name="ALERTAS_MES")
    except Exception:
        return {"status": [], "alerts": [], "source": str(path)}

    status = []
    for _, row in status_df.dropna(how="all").iterrows():
        day = fmt_date(row.get("Dia"))
        if not day:
            continue
        status.append(
            {
                "date": day,
                "hourly": as_number(row.get("Hourly (h)")),
                "daily": row.get("Daily"),
                "sep": row.get("SEP"),
                "status": row.get("Status"),
                "missingHours": row.get("Horas faltando") if pd.notna(row.get("Horas faltando")) else None,
            }
        )
    alerts = []
    for _, row in alert_df.dropna(subset=["day_ref", "issue_type"], how="any").iterrows():
        alerts.append(
            {
                "date": fmt_date(row.get("day_ref")),
                "severity": row.get("severity"),
                "type": row.get("issue_type"),
                "ref": row.get("ref_key"),
                "details": row.get("details"),
            }
        )
    return {"status": status, "alerts": alerts[:50], "source": str(path)}


def _find_checklist_path() -> "Path | None":
    """Encontra o arquivo de Checklist Diário mais recente disponível (mês atual → meses anteriores)."""
    MONTH_PT = {
        1: "jan", 2: "fev", 3: "mar", 4: "abr", 5: "mai", 6: "jun",
        7: "jul", 8: "ago", 9: "set", 10: "out", 11: "nov", 12: "dez",
    }
    MONTH_PT_LONG = {
        1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho",
        7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro",
    }
    rows_to_try: list[str] = []
    for m in range(TODAY.month, 0, -1):
        y2 = str(TODAY.year)[2:]  # 26
        y4 = TODAY.year           # 2026
        mo_short = MONTH_PT[m]
        mo_long  = MONTH_PT_LONG[m]
        rows_to_try += [
            f"Bacalhau - Checklist Diario_{mo_long}_{y4}.xlsm",
            f"Bacalhau - Checklist Diario_{mo_long}_{y2}.xlsm",
            f"Bacalhau - Checklist Diario_{mo_short.capitalize()}-{y2}.xlsm",
            f"Bacalhau - Checklist Diario_{mo_short}-{y2}.xlsm",
        ]
    for name in rows_to_try:
        found = find_first_file(name, "checklistDiario", "physchem", "dailyReports")
        if found and found.exists():
            return found
    return None


def parse_lab_report() -> dict[str, Any]:
    path = _find_checklist_path()
    rows: list[dict[str, Any]] = []
    if path is not None and path.exists():
        try:
            df = pd.read_excel(path, sheet_name="Lab-Report", header=7)
            for _, row in df.iterrows():
                day = fmt_date(row.get("Date"))
                if not day:
                    continue
                rows.append(
                    {
                        "date": day,
                        "labReport": row.get("LabReport"),
                        "api": as_number(row.get("API")),
                        "density": first_number(row, ["Density\n[kg/m3]", "Density [kg/m3]"]),
                        "densityCv": as_number(row.get("Densidade Conf. CV")),
                        "bsw": first_number(row, ["BSW \n[%v/v]", "BSW [%v/v]"]),
                        "bswFlowline": as_number(row.get("BSW Flowline")),
                        "method": row.get("Método"),
                        "source": display_path(path),
                    }
                )
        except Exception:
            rows = []
    rows = [row for row in rows if row.get("api") is not None or row.get("density") is not None or row.get("bsw") is not None]
    rows.sort(key=lambda item: item["date"])
    latest = rows[-1] if rows else {}
    return {
        "source": display_path(path) if path else None,
        "rows": rows[-45:],
        "latest": latest,
        "coverageDays": len({row["date"] for row in rows}),
    }


def clean_model_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and math.isnan(value):
        return ""
    return re.sub(r"\s+", " ", str(value).replace("\x00", " ")).strip()


def parse_model_timestamp(value: Any) -> datetime | None:
    text = clean_model_text(value)
    if not text:
        return None
    parsed = pd.to_datetime(text, errors="coerce", dayfirst=False)
    if pd.isna(parsed):
        parsed = pd.to_datetime(text, errors="coerce", dayfirst=True)
    if pd.isna(parsed):
        return None
    return parsed.to_pydatetime()


def classify_model_domain(file_name: str, source: str = "") -> str:
    text = f"{file_name} {source}".lower()
    if "meteringoverview" in text or "|url" in text:
        return "navigation"
    if "fiscal" in text or "metering bacalhau" in text or "system 20" in text:
        return "fiscal_metering"
    if "gas injection" in text or "injwells" in text or "inj well" in text:
        return "gas_injection"
    if "riser mpfm" in text or "mpm risers" in text:
        return "riser_mpfm"
    if "well mpfm" in text or "mpm wells" in text:
        return "well_mpfm"
    if "flowline" in text:
        return "flowline"
    if "prod well" in text or "prodwellssub" in text:
        return "production_well"
    if "production overview" in text or "quick look" in text or "teste medicao" in text:
        return "production_overview"
    return "operational_model"


def classify_signal_kind(source: str) -> str:
    text = source.lower()
    if "url" in text:
        return "url"
    if any(term in text for term in ("oil", "oleo", "óleo", "gsv", "gsvfr")):
        return "oil"
    if any(term in text for term in ("gas", "gvfr", "flare")):
        return "gas"
    if "water" in text or "bsw" in text:
        return "water"
    if "pressure" in text or "press" in text or "dp " in text or "differential" in text:
        return "pressure"
    if "temperature" in text or "temp" in text:
        return "temperature"
    if "choke" in text:
        return "choke"
    if "valve" in text or "zso" in text or "zsc" in text or "esdv" in text or "xsv" in text:
        return "valve_status"
    if "status" in text or "bad" in text or "mode" in text:
        return "status"
    if "mass" in text:
        return "mass"
    return "other"


def extract_model_asset(source: str) -> str:
    source = clean_model_text(source)
    if not source:
        return "sem_fonte"
    if "|" in source:
        left = source.split("|", 1)[0]
    else:
        left = source
    parts = [part for part in re.split(r"[\\/]", left) if part]
    if parts:
        return parts[-1]
    return source[:80]


def source_label(source: str) -> str:
    source = clean_model_text(source)
    if "|" in source:
        asset, attribute = source.rsplit("|", 1)
        asset_name = extract_model_asset(asset)
        return f"{asset_name}|{attribute}"[:160]
    parts = [part for part in re.split(r"[\\/]", source) if part]
    return (parts[-1] if parts else source)[:160]


def summarize_model_csv_file(path: Path) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    file_summary = {
        "path": display_path(path),
        "name": path.name,
        "domain": classify_model_domain(path.name),
        "rows": 0,
        "numericRows": 0,
        "textRows": 0,
        "badRows": 0,
        "signals": 0,
        "assets": 0,
        "dateStart": None,
        "dateEnd": None,
        "topSignals": [],
        "sampleTextValues": [],
        "warnings": [],
    }
    daily_stats: dict[tuple[str, str, str], dict[str, Any]] = {}
    source_counts: Counter[str] = Counter()
    asset_names: set[str] = set()
    text_samples: list[dict[str, str]] = []
    source_cache: dict[str, tuple[str, str, str]] = {}

    try:
        chunks = pd.read_csv(path, sep=";", dtype=str, chunksize=100000, encoding="utf-8-sig")
    except UnicodeDecodeError:
        chunks = pd.read_csv(path, sep=";", dtype=str, chunksize=100000, encoding="latin1")
    except Exception as exc:
        file_summary["error"] = f"csv: {exc.__class__.__name__}"
        return file_summary, []

    for chunk in chunks:
        if not {"Fonte de dados", "Tempo", "Valor"}.issubset(set(chunk.columns)):
            file_summary["error"] = "colunas obrigatorias ausentes"
            continue
        working = chunk[["Fonte de dados", "Tempo", "Valor"]].copy()
        working.columns = ["source", "timestamp", "value"]
        working["source"] = working["source"].map(clean_model_text)
        working["valueText"] = working["value"].map(clean_model_text)
        working = working[working["source"].astype(bool)]
        if working.empty:
            continue

        timestamp_text = working["timestamp"].map(clean_model_text)
        iso_dates = timestamp_text.str.extract(r"(20\d{2}-\d{2}-\d{2})", expand=False)
        br_dates = timestamp_text.str.extract(r"(\d{2})/(\d{2})/(20\d{2})", expand=True)
        fallback_dates = br_dates[2] + "-" + br_dates[1] + "-" + br_dates[0]
        working["date"] = iso_dates.fillna(fallback_dates)
        bad_time = int(working["date"].isna().sum())
        file_summary["badRows"] += bad_time
        working = working.dropna(subset=["date"])
        if working.empty:
            continue
        file_summary["rows"] += int(len(working))
        chunk_start = str(working["date"].min())
        chunk_end = str(working["date"].max())
        if file_summary["dateStart"] is None or chunk_start < file_summary["dateStart"]:
            file_summary["dateStart"] = chunk_start
        if file_summary["dateEnd"] is None or chunk_end > file_summary["dateEnd"]:
            file_summary["dateEnd"] = chunk_end

        for source, count in working["source"].value_counts().items():
            label = source_label(source)
            source_counts[label] += int(count)
            if source not in source_cache:
                source_cache[source] = (
                    classify_model_domain(path.name, source),
                    classify_signal_kind(source),
                    extract_model_asset(source),
                )
            asset_names.add(source_cache[source][2])

        working["domain"] = working["source"].map(lambda source: source_cache[source][0])
        working["kind"] = working["source"].map(lambda source: source_cache[source][1])
        if file_summary["domain"] == "operational_model" and not working.empty:
            domain_mode = working["domain"].mode()
            if not domain_mode.empty:
                file_summary["domain"] = str(domain_mode.iat[0])

        numeric_text = working["valueText"].str.replace(".", "", regex=False).str.replace(",", ".", regex=False)
        working["numericValue"] = pd.to_numeric(numeric_text, errors="coerce")
        numeric_count = int(working["numericValue"].notna().sum())
        file_summary["numericRows"] += numeric_count
        file_summary["textRows"] += int(len(working) - numeric_count)

        if len(text_samples) < 8:
            text_rows = working[working["numericValue"].isna() & working["valueText"].astype(bool)].head(8 - len(text_samples))
            for _, row in text_rows.iterrows():
                text_samples.append({"source": source_label(row["source"]), "value": row["valueText"][:120], "date": row["date"]})

        numeric = working.dropna(subset=["numericValue"])
        if numeric.empty:
            continue
        grouped = numeric.groupby(["date", "domain", "kind"], dropna=True)["numericValue"].agg(["count", "sum", "min", "max", "last"])
        for (day, domain, kind), row in grouped.iterrows():
            if len(daily_stats) >= MAX_MODEL_FILE_AGGREGATES:
                file_summary["warnings"].append(f"limite de agregados por arquivo atingido: {MAX_MODEL_FILE_AGGREGATES}")
                break
            key = (str(day), str(domain), str(kind))
            stats = daily_stats.setdefault(
                key,
                {"date": key[0], "domain": key[1], "kind": key[2], "count": 0, "sum": 0.0, "min": None, "max": None, "last": None},
            )
            stats["count"] += int(row["count"])
            stats["sum"] += float(row["sum"])
            if not pd.isna(row["min"]):
                row_min = float(row["min"])
                stats["min"] = row_min if stats["min"] is None else min(stats["min"], row_min)
            if not pd.isna(row["max"]):
                row_max = float(row["max"])
                stats["max"] = row_max if stats["max"] is None else max(stats["max"], row_max)
            if not pd.isna(row["last"]):
                stats["last"] = float(row["last"])

    file_summary["signals"] = len(source_counts)
    file_summary["assets"] = len(asset_names)
    if file_summary["badRows"]:
        total_seen = file_summary["rows"] + file_summary["badRows"]
        bad_ratio = file_summary["badRows"] / total_seen if total_seen else 0
        file_summary["warnings"].append(f"{file_summary['badRows']} linhas sem data valida ({bad_ratio:.1%})")
    file_summary["topSignals"] = [{"name": name, "rows": rows} for name, rows in source_counts.most_common(10)]
    file_summary["sampleTextValues"] = text_samples
    aggregates = []
    for item in daily_stats.values():
        count = item["count"]
        aggregates.append(
            {
                "date": item["date"],
                "domain": item["domain"],
                "kind": item["kind"],
                "count": item["count"],
                "avg": item["sum"] / count if count else None,
                "min": item["min"],
                "max": item["max"],
                "last": item["last"],
                "source": display_path(path),
            }
        )
    return file_summary, sorted(aggregates, key=lambda row: (row["date"], row["domain"], row["kind"]))


def parse_measurement_models() -> dict[str, Any]:
    if not MODEL_DIR.exists():
        return {
            "summary": {
                "files": 0,
                "rows": 0,
                "numericRows": 0,
                "signals": 0,
                "dailyAggregatesTotal": 0,
                "dailyAggregatesPublished": 0,
                "outputTruncated": False,
                "warnings": [f"pasta MODELOS nao localizada: {MODEL_DIR}"],
            },
            "files": [],
            "dailyAggregates": [],
        }

    files = []
    daily_aggregates = []
    domain_totals: dict[str, dict[str, Any]] = {}
    warnings: list[str] = []
    for csv_path in sorted(MODEL_DIR.glob("*.csv")):
        file_summary, aggregates = summarize_model_csv_file(csv_path)
        files.append(file_summary)
        for warning in file_summary.get("warnings") or []:
            warnings.append(f"{file_summary.get('name')}: {warning}")
        for row in aggregates:
            try:
                row_date = datetime.fromisoformat(str(row.get("date") or "")).date()
            except ValueError:
                warnings.append(f"{file_summary.get('name')}: agregado com data invalida {row.get('date')}")
                continue
            if REPORTING_PERIOD_START <= row_date <= REPORTING_PERIOD_END:
                daily_aggregates.append(row)
        domain = file_summary.get("domain") or "operational_model"
        domain_summary = domain_totals.setdefault(
            domain,
            {"domain": domain, "files": 0, "rows": 0, "numericRows": 0, "textRows": 0, "signals": 0, "dateStart": None, "dateEnd": None},
        )
        domain_summary["files"] += 1
        domain_summary["rows"] += file_summary.get("rows") or 0
        domain_summary["numericRows"] += file_summary.get("numericRows") or 0
        domain_summary["textRows"] += file_summary.get("textRows") or 0
        domain_summary["signals"] += file_summary.get("signals") or 0
        start = file_summary.get("dateStart")
        end = file_summary.get("dateEnd")
        if start and (domain_summary["dateStart"] is None or start < domain_summary["dateStart"]):
            domain_summary["dateStart"] = start
        if end and (domain_summary["dateEnd"] is None or end > domain_summary["dateEnd"]):
            domain_summary["dateEnd"] = end

    sorted_aggregates = sorted(daily_aggregates, key=lambda row: (row.get("date") or "", row.get("domain") or "", row.get("kind") or ""))
    published_aggregates = sorted_aggregates
    truncated = len(sorted_aggregates) > MAX_MODEL_OUTPUT_AGGREGATES
    if truncated:
        dropped = len(sorted_aggregates) - MAX_MODEL_OUTPUT_AGGREGATES
        dropped_rows = sorted_aggregates[:dropped]
        dropped_dates = [row.get("date") for row in dropped_rows if row.get("date")]
        warnings.append(
            "dailyAggregates limitado para o frontend: "
            f"publicados {MAX_MODEL_OUTPUT_AGGREGATES} de {len(sorted_aggregates)} agregados; "
            f"datas omitidas {min(dropped_dates, default='-')} a {max(dropped_dates, default='-')}"
        )
        published_aggregates = sorted_aggregates[-MAX_MODEL_OUTPUT_AGGREGATES:]

    return {
        "summary": {
            "files": len(files),
            "rows": sum(item.get("rows") or 0 for item in files),
            "numericRows": sum(item.get("numericRows") or 0 for item in files),
            "textRows": sum(item.get("textRows") or 0 for item in files),
            "signals": sum(item.get("signals") or 0 for item in files),
            "dateStart": min((item.get("dateStart") for item in files if item.get("dateStart")), default=None),
            "dateEnd": max((item.get("dateEnd") for item in files if item.get("dateEnd")), default=None),
            "dailyAggregatesTotal": len(sorted_aggregates),
            "dailyAggregatesPublished": len(published_aggregates),
            "outputTruncated": truncated,
            "warnings": warnings[:50],
        },
        "domains": sorted(domain_totals.values(), key=lambda item: item["domain"]),
        "files": files,
        "dailyAggregates": published_aggregates,
    }


def within_range(value: float | None, lower: float | None, upper: float | None) -> str:
    if value is None:
        return "warn"
    if lower is not None and value < lower:
        return "critical"
    if upper is not None and value > upper:
        return "critical"
    return "ok"


def build_limit_monitors(latest_points: list[dict[str, Any]], points: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    monitors = []
    for item in latest_points:
        tag = item["tag"]
        point = points.get(tag, {})
        limits = item.get("limites") or {}
        pressure = limits.get("pressao") or {}
        temp = limits.get("temperatura") or {}
        differential = limits.get("diferencial") or {}
        max_op = point.get("maxOperacao")
        min_op = point.get("minOperacao")
        volume_status = within_range(item.get("volumeCorrigido"), min_op, max_op)
        pressure_status = within_range(item.get("pressao"), pressure.get("lower"), pressure.get("upper"))
        temp_status = within_range(item.get("temperatura"), temp.get("lower"), temp.get("upper"))
        diff_status = within_range(differential.get("value"), differential.get("lower"), differential.get("upper"))
        statuses = [volume_status, pressure_status, temp_status]
        if differential.get("value") is not None:
            statuses.append(diff_status)
        status = "critical" if "critical" in statuses else "warn" if "warn" in statuses else "ok"
        monitors.append(
            {
                "tag": tag,
                "date": item["date"],
                "family": item["family"],
                "familyName": item["familyName"],
                "fluid": point.get("fluid") or item.get("fluid"),
                "meterType": point.get("meterType"),
                "pam": {"lower": min_op, "upper": max_op, "value": item.get("volumeCorrigido"), "status": volume_status},
                "pressure": {"lower": pressure.get("lower"), "upper": pressure.get("upper"), "value": item.get("pressao"), "status": pressure_status},
                "temperature": {"lower": temp.get("lower"), "upper": temp.get("upper"), "value": item.get("temperatura"), "status": temp_status},
                "differential": {
                    "lower": differential.get("lower"),
                    "upper": differential.get("upper"),
                    "value": differential.get("value"),
                    "status": diff_status,
                },
                "alarmEnabled": limits.get("alarmeHabilitado"),
                "status": status,
                "source": item.get("arquivoCarga"),
            }
        )
    return monitors


def build_uncertainty_monitor(comparisons: list[dict[str, Any]], points: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    rows = []
    for item in comparisons:
        point = points.get(item["tag"], {})
        max_unc = point.get("incertezaMaxima")
        rows.append(
            {
                "date": item["date"],
                "tag": item["tag"],
                "family": item["family"],
                "familyName": item["familyName"],
                "fluid": point.get("fluid") or item.get("fluid"),
                "uncertaintyMax": max_unc,
                "dailyUncertainty": None,
                "coverage": "cadastro" if max_unc is not None else "sem cadastro",
                "status": "ok" if max_unc is not None else "warn",
                "source": "Pontos de Medição.xlsx" if max_unc is not None else None,
            }
        )
    return rows


def build_ai_modules(alerts: list[dict[str, Any]]) -> dict[str, Any]:
    critical = [alert for alert in alerts if alert.get("severity") == "critical"]
    return {
        "mode": "copiloto auditavel",
        "principle": "IA explica e correlaciona; regra deterministica decide conformidade.",
        "riskDigest": [
            {
                "severity": "critical" if critical else "warn",
                "title": "Risco regulatorio priorizado",
                "detail": f"{len(critical)} alertas criticos pedem evidencia objetiva e acao documentada.",
            },
            {
                "severity": "warn",
                "title": "Fontes tecnicas configuraveis",
                "detail": "Certificados, incerteza, analises, planos de coleta e PAM ja possuem entradas de configuracao.",
            },
        ],
        "capabilities": [
            {
                "name": "Pergunte ao Radar",
                "detail": "Consultar medicoes, XMLs, falhas, certificados, analises e normas nas pastas configuradas.",
            },
            {
                "name": "Explicar alerta",
                "detail": "Gerar causa provavel, evidencia usada, impacto e recomendacao operacional.",
            },
            {
                "name": "Checklist regulatorio",
                "detail": "Mapear existencia, validade, periodicidade, prazo e base normativa por obrigacao.",
            },
            {
                "name": "Dossie do ponto",
                "detail": "Consolidar cadastro, CV, calibracao, incerteza, analises, historico de falhas, XML e ANP.",
            },
            {
                "name": "Atualizacao assistida",
                "detail": "Ler documentos recebidos e propor atualizacao de limites, faixas, PAM e controles internos.",
            },
        ],
    }


def parse_regulatory_matrix() -> dict[str, Any]:
    if not MATRIX_PATH.exists():
        return {"rows": [], "summary": {"total": 0, "byCategory": [], "bySubcategory": []}}
    rows = json.loads(MATRIX_PATH.read_text(encoding="utf-8"))
    by_category = Counter(row.get("Categoria") for row in rows if row.get("Categoria"))
    by_subcategory = Counter(row.get("Subcategoria") for row in rows if row.get("Subcategoria"))
    return {
        "source": str(MATRIX_PATH),
        "rows": rows,
        "summary": {
            "total": len(rows),
            "byCategory": [{"name": key, "value": value} for key, value in by_category.most_common()],
            "bySubcategory": [{"name": key, "value": value} for key, value in by_subcategory.most_common(12)],
        },
    }


def extract_tags_from_text(text: str) -> list[str]:
    tags = set()
    for pattern in (r"\b\d{2}[A-Z]{2}\d{3,5}[A-Z]?\b", r"\b\d{2}J[NX]\d{3}[A-Z]?\b"):
        tags.update(match.group(0).upper() for match in re.finditer(pattern, text.upper()))
    return sorted(tags)


def infer_date_from_text_path(text: str) -> str | None:
    normalized = text.replace("\\", "/")
    daily = re.search(r"Daily reports_(\d{4}-\d{2}-\d{2})", normalized, re.IGNORECASE)
    if daily:
        return daily.group(1)

    parts = [part for part in normalized.split("/") if part]
    for index, part in enumerate(parts[:-2]):
        if re.fullmatch(r"20\d{2}", part):
            month = re.match(r"(\d{1,2})", parts[index + 1])
            day = re.match(r"(\d{1,2})$", parts[index + 2])
            if month and day:
                try:
                    return date(int(part), int(month.group(1)), int(day.group(1))).isoformat()
                except ValueError:
                    pass
    for index, part in enumerate(parts[:-1]):
        if re.fullmatch(r"20\d{2}", part):
            month = re.match(r"(\d{1,2})", parts[index + 1])
            if month:
                try:
                    return date(int(part), int(month.group(1)), 1).isoformat()
                except ValueError:
                    pass

    file_name = parts[-1] if parts else normalized
    for match in re.finditer(r"(?<!\d)(20\d{2})[-_. ]?(\d{2})[-_. ]?(\d{2})(?!\d)", file_name):
        try:
            return date(int(match.group(1)), int(match.group(2)), int(match.group(3))).isoformat()
        except ValueError:
            pass
    for match in re.finditer(r"(?<!\d)(\d{2})[-_. ](\d{2})[-_. ](20\d{2})(?!\d)", file_name):
        try:
            return date(int(match.group(3)), int(match.group(2)), int(match.group(1))).isoformat()
        except ValueError:
            pass
    return None


def event_datetime_from_line(day_text: str, time_text: str) -> str | None:
    try:
        return datetime.strptime(f"{day_text} {time_text}", "%m/%d/%y %H:%M:%S").isoformat()
    except ValueError:
        return None


def classify_parameter_event(message: str) -> dict[str, Any] | None:
    match = re.search(
        r"Parameter\s+(?P<parameter>.*?)\s+was changed from\s+(?P<old>.*?)\s+to\s+(?P<new>.*?)\s+by\s+(?P<actor>.*)$",
        message,
        re.IGNORECASE,
    )
    if not match:
        return None

    parameter = re.sub(r"\s+", " ", match.group("parameter")).strip()
    lowered = parameter.lower()
    expected = []
    for evidence_type, rule in PARAMETER_RULES.items():
        if any(keyword in lowered for keyword in rule["parameterKeywords"]):
            expected.append(evidence_type)
    if not expected:
        expected = ["pam_limits"]

    return {
        "parameter": parameter,
        "oldValue": match.group("old").strip(),
        "newValue": match.group("new").strip(),
        "actor": match.group("actor").strip(),
        "expectedEvidenceTypes": expected,
        "expectedEvidenceLabels": [PARAMETER_RULES[item]["label"] for item in expected],
    }


def parse_alarm_event_txts(max_files: int = 2500) -> dict[str, Any]:
    files = iter_source_files(("alarmsEvents", "dailyReports"), "*AlarmsAndEvents*.txt")
    files = sorted(files, key=lambda path: ("daily" not in path.name.lower(), str(path)))[:max_files]
    events = []
    seen_events = set()
    alarm_counts = Counter()
    line_count = 0

    for path in files:
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
        except OSError:
            continue
        metadata = {"system": None, "flowComputer": None}
        for line in text.splitlines()[:12]:
            if line.startswith("System"):
                metadata["system"] = line.split(":", 1)[-1].strip()
            elif line.startswith("Flow computer"):
                metadata["flowComputer"] = line.split(":", 1)[-1].strip()

        for line in text.splitlines():
            line_count += 1
            match = re.match(r"\s*(\d{2}/\d{2}/\d{2})\s+(\d{2}:\d{2}:\d{2})\s+(.+?)\s*$", line)
            if not match:
                continue
            message = match.group(3).strip()
            if " alarm " in f" {message.lower()} ":
                label = re.sub(r"\s+from\s+.*$", "", message, flags=re.IGNORECASE)
                alarm_counts[label[:90]] += 1
            parsed = classify_parameter_event(message)
            if not parsed:
                continue
            event_key = (event_datetime_from_line(match.group(1), match.group(2)), message, tuple(extract_tags_from_text(str(path.parent))))
            if event_key in seen_events:
                continue
            seen_events.add(event_key)
            event_text = f"{message} {path.name} {path.parent}"
            events.append(
                {
                    "timestamp": event_key[0],
                    "message": message,
                    "source": display_path(path),
                    "system": metadata["system"],
                    "flowComputer": metadata["flowComputer"],
                    "tags": extract_tags_from_text(event_text),
                    **parsed,
                }
            )

    return {
        "filesScanned": len(files),
        "linesScanned": line_count,
        "events": events,
        "alarmTop": [{"name": name, "value": value} for name, value in alarm_counts.most_common(8)],
    }


def classify_evidence_path(path: Path) -> list[str]:
    text = f"{path.name} {path.parent}".lower()
    types = []
    for evidence_type, rule in PARAMETER_RULES.items():
        if any(keyword in text for keyword in rule["evidenceKeywords"]):
            types.append(evidence_type)
    return types


def index_evidence_files(max_files: int = 12000) -> list[dict[str, Any]]:
    allowed_suffixes = {".pdf", ".xlsx", ".xlsm", ".xls", ".docx", ".doc", ".zip", ".xml", ".txt", ".csv"}
    files: list[Path] = []
    seen = set()
    for base in configured_paths(*EVIDENCE_SOURCE_IDS):
        candidates = [base] if base.is_file() else base.rglob("*")
        for item in candidates:
            if not item.is_file() or is_ignored(item) or item.suffix.lower() not in allowed_suffixes:
                continue
            if "alarmsandevents" in item.name.lower():
                continue
            key = str(item.resolve()).lower()
            if key in seen:
                continue
            seen.add(key)
            files.append(item)
            if len(files) >= max_files:
                break
        if len(files) >= max_files:
            break

    evidence = []
    for path in sorted(files):
        types = classify_evidence_path(path)
        if not types:
            continue
        text = f"{path.name} {path.parent}"
        evidence.append(
            {
                "path": display_path(path),
                "name": path.name,
                "extension": path.suffix.lower(),
                "date": infer_date_from_text_path(str(path)),
                "tags": extract_tags_from_text(text),
                "evidenceTypes": types,
                "evidenceLabels": [PARAMETER_RULES[item]["label"] for item in types],
                "sizeKb": round(path.stat().st_size / 1024, 1),
            }
        )
    return evidence


def resolve_display_file(path_text: str) -> Path:
    path = Path(path_text)
    if path.is_absolute():
        return path
    return ROOT / path


def normalize_evidence_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\x00", " ")).strip()


def extract_pdf_text(path: Path, max_pages: int = 6) -> tuple[str, str | None]:
    try:
        import pdfplumber

        parts = []
        with pdfplumber.open(path) as pdf:
            for index, page in enumerate(pdf.pages[:max_pages], start=1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    parts.append(f"[page {index}] {page_text}")
        return normalize_evidence_text("\n".join(parts)), None
    except Exception as exc:
        return "", f"pdf: {exc.__class__.__name__}"


def extract_workbook_text(path: Path, max_sheets: int = 4, max_rows: int = 250, max_cols: int = 40) -> tuple[str, str | None]:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        parts = []
        for sheet in workbook.worksheets[:max_sheets]:
            parts.append(f"[sheet {sheet.title}]")
            for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                if row_index > max_rows:
                    break
                values = [str(value) for value in row[:max_cols] if value is not None and str(value).strip()]
                if values:
                    parts.append(" | ".join(values))
        workbook.close()
        return normalize_evidence_text("\n".join(parts)), None
    except Exception as exc:
        return "", f"workbook: {exc.__class__.__name__}"


def extract_docx_text(path: Path) -> tuple[str, str | None]:
    try:
        from docx import Document

        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables[:8]:
            for row in table.rows[:80]:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" | ".join(values))
        return normalize_evidence_text("\n".join(parts)), None
    except Exception as exc:
        return "", f"docx: {exc.__class__.__name__}"


def extract_zip_text(path: Path, max_entries: int = 12) -> tuple[str, str | None]:
    text_suffixes = {".txt", ".xml", ".csv", ".html", ".htm"}
    try:
        parts = []
        with zipfile.ZipFile(path) as zf:
            names = zf.namelist()
            parts.append("[zip entries] " + " | ".join(names[:50]))
            read_count = 0
            for name in names:
                suffix = Path(name).suffix.lower()
                if suffix not in text_suffixes:
                    continue
                with zf.open(name) as file:
                    raw = file.read(120000)
                parts.append(f"[zip {name}] " + raw.decode("utf-8", errors="ignore"))
                read_count += 1
                if read_count >= max_entries:
                    break
        return normalize_evidence_text("\n".join(parts)), None
    except Exception as exc:
        return "", f"zip: {exc.__class__.__name__}"


def load_evidence_text_cache() -> dict[str, dict[str, Any]]:
    if not TEXT_CACHE_PATH.exists():
        return {}
    try:
        data = json.loads(TEXT_CACHE_PATH.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    return data if isinstance(data, dict) else {}


def save_evidence_text_cache(cache: dict[str, dict[str, Any]]) -> None:
    TEXT_CACHE_PATH.parent.mkdir(parents=True, exist_ok=True)
    TEXT_CACHE_PATH.write_text(json.dumps(cache, ensure_ascii=False), encoding="utf-8")


def extract_evidence_text(evidence: dict[str, Any], cache: dict[str, dict[str, Any]]) -> dict[str, Any]:
    path = resolve_display_file(evidence["path"])
    if not path.exists():
        return {"text": "", "error": "arquivo nao localizado", "extractor": None}
    stat = path.stat()
    key = f"{str(path.resolve()).lower()}::{stat.st_mtime_ns}::{stat.st_size}"
    if key in cache:
        return cache[key]

    suffix = path.suffix.lower()
    text = ""
    error = None
    extractor = suffix.lstrip(".")
    if suffix == ".pdf":
        text, error = extract_pdf_text(path)
    elif suffix in {".xlsx", ".xlsm"}:
        text, error = extract_workbook_text(path)
    elif suffix == ".docx":
        text, error = extract_docx_text(path)
    elif suffix == ".zip":
        text, error = extract_zip_text(path)
    elif suffix in {".txt", ".xml", ".csv", ".html", ".htm"}:
        try:
            text = normalize_evidence_text(path.read_text(encoding="utf-8", errors="ignore")[:180000])
        except OSError as exc:
            error = f"text: {exc.__class__.__name__}"
    else:
        error = "formato ainda nao extraido"

    result = {"text": text[:80000], "error": error, "extractor": extractor, "chars": len(text)}
    cache[key] = result
    return result


def parameter_terms(parameter: str) -> list[str]:
    ignored = {"run", "parameter", "override", "was", "changed", "from", "to", "by"}
    terms = []
    for term in re.findall(r"[a-zA-Z][a-zA-Z0-9+-]*", parameter.lower()):
        if len(term) >= 3 and term not in ignored:
            terms.append(term)
    if "carbon" in terms and "dioxide" in terms:
        terms.append("carbon dioxide")
    if "carbon" in terms and "monoxide" in terms:
        terms.append("carbon monoxide")
    if "dynamic" in terms and "viscosity" in terms:
        terms.append("dynamic viscosity")
    if "isentropic" in terms and "exponent" in terms:
        terms.append("isentropic exponent")
    return sorted(set(terms), key=len, reverse=True)


def numeric_variants(value: str) -> list[str]:
    variants = {value.strip().lower()}
    match = re.search(r"[-+]?\d+(?:\.\d+)?(?:e[-+]?\d+)?", value.strip(), re.IGNORECASE)
    if match:
        try:
            number_value = float(match.group(0))
            variants.add(f"{number_value:.12g}".lower())
            variants.add(f"{number_value:.6g}".lower())
            variants.add(f"{number_value:.5f}".rstrip("0").rstrip(".").lower())
        except ValueError:
            pass
    return sorted(item for item in variants if item)


def find_content_snippet(text: str, terms: list[str], values: list[str]) -> tuple[str | None, list[str], int | None]:
    lowered = text.lower()
    hits = []
    term_positions = []
    value_positions = []
    for term in terms:
        pos = lowered.find(term)
        if pos >= 0:
            hits.append(f"parametro:{term}")
            term_positions.append(pos)
    for value in values:
        if not value:
            continue
        pos = lowered.find(value)
        if pos >= 0:
            hits.append(f"valor:{value}")
            value_positions.append(pos)

    distance = None
    if term_positions and value_positions:
        distance = min(abs(left - right) for left in term_positions for right in value_positions)
        anchor = min(
            ((abs(left - right), min(left, right), max(left, right)) for left in term_positions for right in value_positions),
            key=lambda item: item[0],
        )
        start = max(0, anchor[1] - 140)
        end = min(len(text), anchor[2] + 280)
    elif term_positions:
        start = max(0, term_positions[0] - 140)
        end = min(len(text), term_positions[0] + 360)
    elif value_positions:
        start = max(0, value_positions[0] - 140)
        end = min(len(text), value_positions[0] + 360)
    else:
        return None, hits, distance
    snippet = text[start:end].strip()
    return snippet, hits, distance


def evaluate_content_evidence(event: dict[str, Any], evidence: dict[str, Any], cache: dict[str, dict[str, Any]]) -> dict[str, Any]:
    extraction = extract_evidence_text(evidence, cache)
    text = extraction.get("text") or ""
    terms = parameter_terms(event.get("parameter") or "")
    new_variants = numeric_variants(event.get("newValue") or "")
    snippet, hits, hit_distance = find_content_snippet(text, terms, new_variants)

    state = "candidate"
    reason = "classificado por caminho/nome"
    if extraction.get("error"):
        state = "candidate"
        reason = f"conteudo nao extraido: {extraction['error']}"
    elif "parametro" in " ".join(hits) and "valor" in " ".join(hits) and hit_distance is not None and hit_distance <= 1200:
        state = "confirmed"
        reason = "conteudo contem parametro e valor novo proximos"
    elif "parametro" in " ".join(hits) and "valor" in " ".join(hits):
        state = "supporting"
        reason = "conteudo contem parametro e valor novo em pontos distantes"
    elif "parametro" in " ".join(hits):
        state = "supporting"
        reason = "conteudo contem parametro, sem confirmar valor novo"
    elif text:
        type_terms = []
        for evidence_type in event.get("expectedEvidenceTypes") or []:
            type_terms.extend(PARAMETER_RULES[evidence_type]["evidenceKeywords"])
        type_hit = next((term for term in type_terms if term and term.lower() in text.lower()), None)
        if type_hit:
            state = "supporting"
            reason = f"conteudo contem termo documental: {type_hit}"

    return {
        "contentState": state,
        "contentReason": reason,
        "contentHits": hits,
        "contentDistance": hit_distance,
        "snippet": snippet,
        "extractor": extraction.get("extractor"),
        "contentChars": extraction.get("chars") or 0,
    }


def date_distance_days(left: str | None, right: str | None) -> int | None:
    if not left or not right:
        return None
    try:
        return abs((datetime.fromisoformat(left[:10]).date() - datetime.fromisoformat(right[:10]).date()).days)
    except ValueError:
        return None


def score_evidence_match(event: dict[str, Any], evidence: dict[str, Any]) -> tuple[int, list[str]]:
    score = 0
    reasons = []
    expected = set(event.get("expectedEvidenceTypes") or [])
    evidence_types = set(evidence.get("evidenceTypes") or [])
    overlap_types = sorted(expected & evidence_types)
    if overlap_types:
        score += 3
        reasons.append("tipo esperado")
    else:
        return 0, []

    event_tags = set(event.get("tags") or [])
    evidence_tags = set(evidence.get("tags") or [])
    if event_tags and evidence_tags and event_tags & evidence_tags:
        score += 4
        reasons.append("tag/equipamento")

    distance = date_distance_days(event.get("timestamp"), evidence.get("date"))
    if distance is not None:
        if distance <= 7:
            score += 3
            reasons.append("data proxima")
        elif distance <= 45:
            score += 2
            reasons.append("mesma janela")
        elif distance <= 370:
            score += 1
            reasons.append("mesmo ciclo anual")

    source_text = f"{event.get('system') or ''} {event.get('flowComputer') or ''}".lower()
    path_text = f"{evidence.get('name')} {evidence.get('path')}".lower()
    if source_text and any(piece and piece in path_text for piece in re.split(r"\s+-\s+|\s+", source_text) if len(piece) >= 5):
        score += 1
        reasons.append("contexto do sistema")
    return score, reasons


def build_event_evidence_radar() -> dict[str, Any]:
    parsed_events = parse_alarm_event_txts()
    evidence_index = index_evidence_files()
    matched = []
    status_counts = Counter()
    content_counts = Counter()
    extraction_cache = load_evidence_text_cache()

    for event in parsed_events["events"]:
        candidates = []
        for evidence in evidence_index:
            score, reasons = score_evidence_match(event, evidence)
            if score:
                candidates.append({**evidence, "score": score, "reasons": reasons})
        candidates = sorted(candidates, key=lambda item: (-item["score"], item.get("date") or "", item["path"]))[:6]
        enriched_candidates = []
        for candidate in candidates:
            content = evaluate_content_evidence(event, candidate, extraction_cache)
            content_bonus = 4 if content["contentState"] == "confirmed" else 2 if content["contentState"] == "supporting" else 0
            enriched_candidates.append(
                {
                    **candidate,
                    **content,
                    "score": candidate["score"] + content_bonus,
                    "baseScore": candidate["score"],
                }
            )
        enriched_candidates = sorted(
            enriched_candidates,
            key=lambda item: (
                {"confirmed": 0, "supporting": 1, "candidate": 2}.get(item.get("contentState"), 9),
                -item["score"],
                item.get("date") or "",
                item["path"],
            ),
        )[:5]
        best_score = enriched_candidates[0]["score"] if enriched_candidates else 0
        best_state = enriched_candidates[0].get("contentState") if enriched_candidates else None
        if best_state == "confirmed" and best_score >= 7:
            status = "ok"
        elif enriched_candidates:
            status = "warn"
        else:
            status = "critical"
        status_counts[status] += 1
        content_counts[best_state or "missing"] += 1
        matched.append(
            {
                **event,
                "status": status,
                "bestScore": best_score,
                "evidenceState": best_state or "missing",
                "evidenceCandidates": enriched_candidates,
            }
        )

    matched.sort(key=lambda item: ({"critical": 0, "warn": 1, "ok": 2}.get(item["status"], 9), item.get("timestamp") or ""))
    save_evidence_text_cache(extraction_cache)
    evidence_by_type = Counter(evidence_type for item in evidence_index for evidence_type in item.get("evidenceTypes", []))
    return {
        "summary": {
            "eventFilesScanned": parsed_events["filesScanned"],
            "eventLinesScanned": parsed_events["linesScanned"],
            "parameterChanges": len(parsed_events["events"]),
            "evidenceIndexed": len(evidence_index),
            "ok": status_counts.get("ok", 0),
            "warn": status_counts.get("warn", 0),
            "critical": status_counts.get("critical", 0),
            "confirmed": content_counts.get("confirmed", 0),
            "supporting": content_counts.get("supporting", 0),
            "candidateOnly": content_counts.get("candidate", 0),
        },
        "rules": [
            {"id": key, "label": value["label"], "keywords": value["parameterKeywords"]}
            for key, value in PARAMETER_RULES.items()
        ],
        "events": matched[:120],
        "alarmTop": parsed_events["alarmTop"],
        "evidenceByType": [
            {"name": PARAMETER_RULES[key]["label"], "value": value}
            for key, value in evidence_by_type.most_common()
        ],
    }


def compare_layers(
    xml_records: list[dict[str, Any]],
    anp_index: dict[tuple[str, str, str], dict[str, Any]],
    cv_index: dict[tuple[str, str], dict[str, Any]],
) -> list[dict[str, Any]]:
    rows = []
    for xml in xml_records:
        family = xml["family"]
        if family not in {"a001", "a002", "a003"}:
            continue
        day = xml.get("date")
        tag = xml.get("tag")
        anp = anp_index.get((day, family, tag))
        cv = cv_index.get((day, tag))
        scale = 1000 if family in {"a002", "a003"} else 1
        xml_corr = xml.get("volumeCorrigido")
        cv_corr = cv.get("volumeCorrigido") if cv else None
        cv_comp = cv_corr / scale if cv_corr is not None else None
        anp_corr = anp.get("volumeCorrigido") if anp else None

        def close(a: float | None, b: float | None, tol: float = 0.02) -> bool:
            return a is not None and b is not None and abs(a - b) <= tol

        raw_ok = close(xml_corr, cv_comp)
        anp_ok = close(xml_corr, anp_corr)
        status = "ok" if raw_ok and anp_ok else "warn" if anp_ok else "critical"
        if not cv:
            status = "warn"
        rows.append(
            {
                "date": day,
                "tag": tag,
                "family": family,
                "familyName": FAMILIES[family]["name"],
                "fluid": FAMILIES[family]["fluid"],
                "rawCorrigido": cv_comp,
                "xmlCorrigido": xml_corr,
                "anpCorrigido": anp_corr,
                "rawSource": cv.get("source") if cv else None,
                "xmlSource": xml.get("source"),
                "rawOk": raw_ok,
                "anpOk": anp_ok,
                "status": status,
                "note": "Sem Run_Daily CV localizado" if not cv else None,
            }
        )
    return sorted(rows, key=lambda x: (x["date"] or "", x["family"], x["tag"]))


def build_daily_closing(rows: list[dict[str, Any]], files: list[dict[str, Any]], bsw_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    by_date: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in rows:
        by_date[row["date"]].append(row)
    bsw_by_date = defaultdict(list)
    for row in bsw_rows:
        if row.get("date"):
            bsw_by_date[row["date"]].append(row)
    file_by_date = defaultdict(list)
    for item in files:
        if item.get("date"):
            file_by_date[item["date"]].append(item)

    closing = []
    expected = {"a001", "a002", "a003", "a004"}
    for day, day_rows in sorted(by_date.items()):
        present = {item["family"] for item in file_by_date[day]}
        missing = sorted(expected - present)
        total_oil = sum((row.get("anpCorrigido") or 0) for row in day_rows if row["family"] == "a001")
        total_gas = sum((row.get("anpCorrigido") or 0) for row in day_rows if row["family"] in {"a002", "a003"})
        warn = sum(1 for row in day_rows if row["status"] == "warn")
        critical = sum(1 for row in day_rows if row["status"] == "critical")
        closing.append(
            {
                "date": day,
                "status": "critical" if critical or missing else "warn" if warn else "ok",
                "points": len(day_rows),
                "ok": sum(1 for row in day_rows if row["status"] == "ok"),
                "warn": warn,
                "critical": critical,
                "missingFamilies": missing,
                "totalOil": total_oil,
                "totalGas": total_gas,
                "maxBsw": max((row.get("bsw") or 0) for row in bsw_by_date[day]) if bsw_by_date[day] else None,
            }
        )
    return closing


def build_alerts(
    closing: list[dict[str, Any]],
    rows: list[dict[str, Any]],
    failures: dict[str, Any],
    mpfm: dict[str, Any],
    operator_panel_health: dict[str, Any],
) -> list[dict[str, Any]]:
    alerts = []
    if operator_panel_health.get("status") != "ok":
        for item in operator_panel_health.get("exports", []):
            if item.get("status") == "ok":
                continue
            alerts.append(
                {
                    "severity": "critical",
                    "date": None,
                    "title": f"Painel do Operador incompleto: {item.get('label')}",
                    "detail": item.get("message"),
                    "area": "Painel do Operador",
                }
            )
    for day in closing:
        if day["missingFamilies"]:
            alerts.append(
                {
                    "severity": "critical",
                    "date": day["date"],
                    "title": "XML esperado sem evidencia",
                    "detail": "Familias ausentes: " + ", ".join(day["missingFamilies"]),
                    "area": "Envio ANP",
                }
            )
    for row in rows:
        if row["status"] != "ok":
            alerts.append(
                {
                    "severity": "warn" if row["status"] == "warn" else "critical",
                    "date": row["date"],
                    "title": f"Trilha incompleta {row['tag']}",
                    "detail": row.get("note") or "Divergencia entre raw, XML e Painel ANP",
                    "area": row["familyName"],
                }
            )
    for item in failures.get("latestOpen", [])[:5]:
        if (item.get("overdueDays") or 0) > 0:
            alerts.append(
                {
                    "severity": "critical",
                    "date": item.get("detected"),
                    "title": f"Falha aberta {item.get('code')}",
                    "detail": f"{item.get('tag')} com previsao vencida ha {item.get('overdueDays')} dias",
                    "area": "NFSM",
                }
            )
    for item in mpfm.get("alerts", [])[:8]:
        alerts.append(
            {
                "severity": "warn",
                "date": item.get("date"),
                "title": item.get("type"),
                "detail": item.get("details"),
                "area": "MPFM",
            }
        )
    severity_rank = {"critical": 0, "warn": 1, "ok": 2}
    alerts.sort(key=lambda x: (severity_rank.get(x.get("severity"), 9), x.get("date") or ""))
    return alerts[:18]


def stable_id(prefix: str, *parts: Any) -> str:
    identity = json.dumps(parts, ensure_ascii=False, sort_keys=True, default=str)
    return f"{prefix}-{hashlib.sha1(identity.encode('utf-8')).hexdigest()[:10].upper()}"


def load_pendency_decisions() -> dict[str, dict[str, Any]]:
    try:
        raw = json.loads(PENDENCY_DECISIONS_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}
    if isinstance(raw, dict) and isinstance(raw.get("decisions"), dict):
        return raw["decisions"]
    return raw if isinstance(raw, dict) else {}


def unique_pending_items(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    seen: set[str] = set()
    unique = []
    for item in items:
        item_id = item.get("id")
        if not item_id or item_id in seen:
            continue
        seen.add(item_id)
        unique.append(item)
    return unique


def build_operational_calendar(
    dates: list[str],
    files: list[dict[str, Any]],
    comparisons: list[dict[str, Any]],
    closing: list[dict[str, Any]],
    alerts: list[dict[str, Any]],
) -> dict[str, Any]:
    decisions = load_pendency_decisions()
    loaded_dates = set(dates)
    file_by_date: dict[str, list[dict[str, Any]]] = defaultdict(list)
    rows_by_date: dict[str, list[dict[str, Any]]] = defaultdict(list)
    alerts_by_date: dict[str, list[dict[str, Any]]] = defaultdict(list)
    closing_by_date = {row.get("date"): row for row in closing}

    for item in files:
        if item.get("date"):
            file_by_date[item["date"]].append(item)
    for item in comparisons:
        if item.get("date"):
            rows_by_date[item["date"]].append(item)
    for item in alerts:
        if item.get("date"):
            alerts_by_date[item["date"]].append(item)

    if dates:
        first = datetime.fromisoformat(min(dates)).date()
        month_start = min(REPORTING_PERIOD_START, first.replace(day=1))
        last = max(REPORTING_PERIOD_END, TODAY, datetime.fromisoformat(max(dates)).date())
    else:
        month_start = REPORTING_PERIOD_START
        last = max(REPORTING_PERIOD_END, TODAY)

    day_count = (last - month_start).days + 1
    days = []
    expected_families = {"a001", "a002", "a003", "a004"}

    for offset in range(day_count):
        current = month_start + timedelta(days=offset)
        day = current.isoformat()
        day_files = file_by_date.get(day, [])
        day_rows = rows_by_date.get(day, [])
        day_alerts = alerts_by_date.get(day, [])
        closing_row = closing_by_date.get(day)
        xml_families = sorted({item.get("family") for item in day_files if item.get("kind") == "xml" and item.get("family")})
        package_families = sorted({item.get("family") for item in day_files if item.get("kind") == "zip" and item.get("family")})
        pending_items = []

        if day not in loaded_dates:
            pending_items.append(
                {
                    "id": stable_id("PEND", day, "source", "not_loaded"),
                    "type": "missing_source",
                    "severity": "critical" if current < TODAY else "warn",
                    "title": "Dia sem dados carregados",
                    "detail": "Nenhum raw/XML/Painel ANP foi encontrado para esta data.",
                    "recommendedAction": "Corrigir caminho/fonte ou carregar os arquivos do dia e reprocessar.",
                    "resolutionMode": "corrigir_fonte",
                }
            )

        missing_xml = sorted(expected_families - set(xml_families))
        if day in loaded_dates and missing_xml:
            pending_items.append(
                {
                    "id": stable_id("PEND", day, "xml", missing_xml),
                    "type": "missing_xml_family",
                    "severity": "warn",
                    "title": "Família XML sem arquivo extraído",
                    "detail": "Famílias sem XML direto: " + ", ".join(missing_xml),
                    "recommendedAction": "Conferir se a família veio compactada, exportar XML direto ou aceitar justificativa operacional.",
                    "resolutionMode": "corrigir_fonte_ou_baixa",
                }
            )

        raw_pending = sum(1 for row in day_rows if not row.get("rawOk"))
        anp_pending = sum(1 for row in day_rows if not row.get("anpOk"))
        if raw_pending:
            pending_items.append(
                {
                    "id": stable_id("PEND", day, "raw", raw_pending),
                    "type": "raw_xml_mismatch",
                    "severity": "warn",
                    "title": "Raw x XML com pendência",
                    "detail": f"{raw_pending} ponto(s) sem conciliação raw -> XML.",
                    "recommendedAction": "Corrigir raw/Run_Daily na fonte ou registrar baixa com justificativa.",
                    "resolutionMode": "corrigir_fonte_ou_baixa",
                }
            )
        if anp_pending:
            pending_items.append(
                {
                    "id": stable_id("PEND", day, "anp", anp_pending),
                    "type": "xml_anp_mismatch",
                    "severity": "critical",
                    "title": "XML x Painel ANP com pendência",
                    "detail": f"{anp_pending} ponto(s) sem conciliação XML -> ANP.",
                    "recommendedAction": "Verificar envio/recebimento ANP e reprocessar após correção.",
                    "resolutionMode": "corrigir_fonte",
                }
            )

        for alert in day_alerts[:6]:
            pending_items.append(
                {
                    "id": stable_id("PEND", day, "alert", alert.get("title"), alert.get("detail"), alert.get("area")),
                    "type": "alert",
                    "severity": alert.get("severity") or "warn",
                    "title": alert.get("title"),
                    "detail": alert.get("detail"),
                    "recommendedAction": "Tratar causa, corrigir fonte quando aplicável ou registrar baixa operacional justificada.",
                    "resolutionMode": "corrigir_fonte_ou_baixa",
                }
            )

        pending_items = unique_pending_items(pending_items)

        for item in pending_items:
            decision = decisions.get(item["id"])
            if decision:
                item["status"] = decision.get("decision")
                item["closedBy"] = decision.get("closedBy")
                item["closedAt"] = decision.get("closedAt")
                item["decisionNote"] = decision.get("note")
            else:
                item["status"] = "open"

        open_items = [item for item in pending_items if item.get("status") == "open"]
        if not day_rows and day not in loaded_dates:
            status = "not_loaded"
        elif any(item.get("severity") == "critical" for item in open_items):
            status = "critical"
        elif open_items or (closing_row and closing_row.get("status") != "ok"):
            status = "warn"
        elif pending_items:
            status = "resolved"
        else:
            status = "ok"

        days.append(
            {
                "date": day,
                "day": current.day,
                "loaded": day in loaded_dates,
                "status": status,
                "closingStatus": closing_row.get("status") if closing_row else None,
                "points": len(day_rows),
                "xmlFamilies": xml_families,
                "packageFamilies": package_families,
                "missingXmlFamilies": missing_xml if day in loaded_dates else [],
                "rawPending": raw_pending,
                "anpPending": anp_pending,
                "alertCount": len(day_alerts),
                "pendingCount": len(pending_items),
                "openPendingCount": len(open_items),
                "resolvedPendingCount": len(pending_items) - len(open_items),
                "pendingItems": pending_items,
            }
        )

    counts = Counter(day["status"] for day in days)
    return {
        "month": f"{month_start:%Y-%m} a {last:%Y-%m}",
        "start": month_start.isoformat(),
        "end": last.isoformat(),
        "summary": {
            "days": len(days),
            "loaded": sum(1 for day in days if day["loaded"]),
            "ok": counts.get("ok", 0),
            "warn": counts.get("warn", 0),
            "critical": counts.get("critical", 0),
            "notLoaded": counts.get("not_loaded", 0),
            "resolved": counts.get("resolved", 0),
            "openPendencies": sum(day["openPendingCount"] for day in days),
        },
        "days": sorted(days, key=lambda item: item["date"]),
    }


def load_proposal_decisions() -> dict[str, dict[str, Any]]:
    try:
        raw = json.loads(CANDIDATE_DECISIONS_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}
    if isinstance(raw, dict) and isinstance(raw.get("decisions"), dict):
        return raw["decisions"]
    return raw if isinstance(raw, dict) else {}


def proposal_confidence(evidence_state: str | None, score: float | int | None = None) -> str:
    if evidence_state == "confirmed" and (score or 0) >= 7:
        return "alta"
    if evidence_state in {"confirmed", "supporting"}:
        return "media"
    return "baixa"


def proposal_risk(parameter: str | None, status: str | None = None) -> str:
    text = (parameter or "").lower()
    if status == "critical" or any(term in text for term in ["dens", "density", "bsw", "chrom", "crom", "pvt", "factor", "fator"]):
        return "alto"
    if any(term in text for term in ["range", "limit", "limite", "pressure", "press", "temp"]):
        return "medio"
    return "baixo"


def apply_proposal_decisions(proposals: list[dict[str, Any]]) -> list[dict[str, Any]]:
    decisions = load_proposal_decisions()
    for proposal in proposals:
        decision = decisions.get(proposal["id"])
        if not decision:
            continue
        proposal["status"] = decision.get("decision") or proposal["status"]
        proposal["authorizedBy"] = decision.get("authorizedBy")
        proposal["authorizedAt"] = decision.get("authorizedAt")
        proposal["decisionNote"] = decision.get("note")
        proposal.setdefault("auditTrail", []).append(
            {
                "at": decision.get("authorizedAt"),
                "actor": decision.get("authorizedBy") or "usuario",
                "action": decision.get("decision"),
                "note": decision.get("note"),
                "source": "proposal-decisions.json",
            }
        )
    return proposals


def build_change_proposals(
    limit_monitors: list[dict[str, Any]],
    uncertainty_monitor: list[dict[str, Any]],
    regulatory_matrix: dict[str, Any],
    event_evidence_radar: dict[str, Any],
) -> list[dict[str, Any]]:
    proposals: list[dict[str, Any]] = []
    proposal_ids: set[str] = set()

    def add(item: dict[str, Any]) -> None:
        identity = json.dumps(
            [
                item.get("kind"),
                item.get("targetType"),
                item.get("targetId"),
                item.get("field"),
                item.get("currentValue"),
                item.get("proposedValue"),
                item.get("sourcePath"),
            ],
            ensure_ascii=False,
            sort_keys=True,
            default=str,
        )
        base_id = f"PROP-{hashlib.sha1(identity.encode('utf-8')).hexdigest()[:10].upper()}"
        proposal_id = base_id
        collision = 2
        while proposal_id in proposal_ids:
            proposal_id = f"{base_id}-{collision}"
            collision += 1
        proposal_ids.add(proposal_id)
        item["id"] = proposal_id
        item.setdefault("status", "pending_authorization")
        item.setdefault("requiresApproval", True)
        item.setdefault("createdAt", datetime.now().isoformat(timespec="seconds"))
        item.setdefault("auditTrail", []).insert(
            0,
            {
                "at": item["createdAt"],
                "actor": "radar-anp",
                "action": "proposal_created",
                "note": "Achado criado a partir dos documentos e planilhas ja analisados.",
            },
        )
        proposals.append(clean_json(item))

    for event in (event_evidence_radar.get("events") or [])[:35]:
        if not event.get("newValue"):
            continue
        best = (event.get("evidenceCandidates") or [{}])[0] or {}
        tags = event.get("tags") or []
        target = ", ".join(tags) if tags else event.get("flowComputer") or event.get("system") or "sistema"
        add(
            {
                "kind": "parameter_update",
                "domain": "Eventos e parametros",
                "title": f"Atualizar {event.get('parameter') or 'parametro'} em {target}",
                "targetType": "ponto_ou_computador_vazao",
                "targetId": target,
                "field": event.get("parameter"),
                "currentValue": event.get("oldValue"),
                "proposedValue": event.get("newValue"),
                "unit": None,
                "confidence": proposal_confidence(event.get("evidenceState"), event.get("bestScore")),
                "risk": proposal_risk(event.get("parameter"), event.get("status")),
                "sourceType": "alarm_event_with_document_evidence",
                "sourcePath": best.get("path"),
                "sourceName": best.get("name") or event.get("source"),
                "evidenceState": event.get("evidenceState") or "missing",
                "evidenceText": best.get("snippet") or best.get("contentReason") or "; ".join(best.get("reasons") or []),
                "recommendedAction": "Autorizar somente se a evidencia confirmar que o novo valor deve atualizar cadastro/controle interno.",
            }
        )

    seen_uncertainty: set[str] = set()
    for row in uncertainty_monitor:
        tag = row.get("tag")
        if not tag or tag in seen_uncertainty or row.get("uncertaintyMax") is None:
            continue
        seen_uncertainty.add(tag)
        add(
            {
                "kind": "dossier_update",
                "domain": "Incerteza",
                "title": f"Registrar incerteza maxima no dossie de {tag}",
                "targetType": "ponto_medicao",
                "targetId": tag,
                "field": "uncertaintyMax",
                "currentValue": None,
                "proposedValue": row.get("uncertaintyMax"),
                "unit": "%",
                "confidence": "media",
                "risk": "medio",
                "sourceType": "cadastro_anp",
                "sourcePath": row.get("source"),
                "sourceName": row.get("source"),
                "evidenceState": "supporting",
                "evidenceText": f"Valor identificado no cadastro: {row.get('uncertaintyMax')}%",
                "recommendedAction": "Conferir memoria de calculo/certificado antes de promover para cadastro mestre.",
            }
        )

    for monitor in limit_monitors:
        if monitor.get("status") == "ok":
            continue
        pam = monitor.get("pam") or {}
        add(
            {
                "kind": "limit_review",
                "domain": "Limites/PAM",
                "title": f"Revisar envelope operacional de {monitor.get('tag')}",
                "targetType": "ponto_medicao",
                "targetId": monitor.get("tag"),
                "field": "pam",
                "currentValue": pam.get("value"),
                "proposedValue": {"lower": pam.get("lower"), "upper": pam.get("upper")},
                "unit": monitor.get("fluid"),
                "confidence": "media",
                "risk": "alto" if monitor.get("status") == "critical" else "medio",
                "sourceType": "painel_anp_vs_cadastro",
                "sourcePath": monitor.get("source"),
                "sourceName": monitor.get("source"),
                "evidenceState": "supporting",
                "evidenceText": f"Valor {pam.get('value')} comparado ao envelope {pam.get('lower')} - {pam.get('upper')}.",
                "recommendedAction": "Nao atualizar limite automaticamente; abrir revisao tecnica ou registrar justificativa operacional.",
            }
        )

    for row in (regulatory_matrix.get("rows") or [])[:12]:
        requisito = row.get("Requisito") or row.get("Item") or row.get("Obrigacao")
        if not requisito:
            continue
        add(
            {
                "kind": "regulatory_control",
                "domain": "Checklist regulatorio",
                "title": str(requisito)[:120],
                "targetType": "obrigacao_regulatoria",
                "targetId": row.get("Subcategoria") or row.get("Categoria") or "SGM",
                "field": "periodicidade_prazo",
                "currentValue": None,
                "proposedValue": {"periodicidade": row.get("Periodicidade"), "prazo": row.get("Prazo")},
                "unit": None,
                "confidence": "media",
                "risk": "medio",
                "sourceType": "matriz_requisitos",
                "sourcePath": regulatory_matrix.get("source"),
                "sourceName": "Matriz de requisitos metrologicos",
                "evidenceState": "supporting",
                "evidenceText": f"Periodicidade: {row.get('Periodicidade') or '-'}; prazo: {row.get('Prazo') or '-'}.",
                "recommendedAction": "Autorizar para entrar no checklist vivo de obrigacoes e prazos.",
            }
        )

    severity_order = {"alto": 0, "medio": 1, "baixo": 2}
    proposals.sort(key=lambda item: (severity_order.get(item.get("risk"), 9), item.get("confidence") != "alta", item.get("title") or ""))
    return apply_proposal_decisions(proposals[:75])


def json_text(value: Any) -> str:
    return json.dumps(clean_json(value), ensure_ascii=False, allow_nan=False)


def bool_int(value: Any) -> int:
    return 1 if bool(value) else 0


def write_sqlite_database(data: dict[str, Any]) -> dict[str, Any]:
    SQLITE_PATH.parent.mkdir(parents=True, exist_ok=True)
    if SQLITE_PATH.exists():
        SQLITE_PATH.unlink()

    conn = sqlite3.connect(SQLITE_PATH)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    cur = conn.cursor()

    cur.executescript(
        """
        CREATE TABLE metadata (
            key TEXT PRIMARY KEY,
            value TEXT
        );

        CREATE TABLE kpis (
            metric TEXT PRIMARY KEY,
            value TEXT
        );

        CREATE TABLE sources (
            id TEXT PRIMARY KEY,
            label TEXT,
            kind TEXT,
            recursive INTEGER,
            description TEXT,
            paths_json TEXT
        );

        CREATE TABLE files (
            id INTEGER PRIMARY KEY,
            date TEXT,
            family TEXT,
            family_name TEXT,
            kind TEXT,
            path TEXT,
            records INTEGER,
            row_json TEXT
        );

        CREATE TABLE comparisons (
            id INTEGER PRIMARY KEY,
            date TEXT,
            tag TEXT,
            family TEXT,
            family_name TEXT,
            fluid TEXT,
            raw_corrigido REAL,
            xml_corrigido REAL,
            anp_corrigido REAL,
            raw_ok INTEGER,
            anp_ok INTEGER,
            status TEXT,
            note TEXT,
            raw_source TEXT,
            xml_source TEXT,
            row_json TEXT
        );

        CREATE TABLE closing (
            id INTEGER PRIMARY KEY,
            date TEXT,
            status TEXT,
            points INTEGER,
            ok INTEGER,
            warn INTEGER,
            critical INTEGER,
            total_oil REAL,
            total_gas REAL,
            max_bsw REAL,
            missing_families_json TEXT,
            row_json TEXT
        );

        CREATE TABLE alerts (
            id INTEGER PRIMARY KEY,
            severity TEXT,
            date TEXT,
            title TEXT,
            detail TEXT,
            area TEXT,
            row_json TEXT
        );

        CREATE TABLE limit_monitors (
            id INTEGER PRIMARY KEY,
            tag TEXT,
            family_name TEXT,
            fluid TEXT,
            current_value REAL,
            min_operacao REAL,
            max_operacao REAL,
            pam REAL,
            status TEXT,
            row_json TEXT
        );

        CREATE TABLE uncertainty_monitor (
            id INTEGER PRIMARY KEY,
            date TEXT,
            tag TEXT,
            family TEXT,
            family_name TEXT,
            uncertainty_max REAL,
            daily_uncertainty REAL,
            coverage TEXT,
            status TEXT,
            source TEXT,
            row_json TEXT
        );

        CREATE TABLE event_evidence_events (
            id INTEGER PRIMARY KEY,
            timestamp TEXT,
            system TEXT,
            flow_computer TEXT,
            tags_json TEXT,
            parameter TEXT,
            old_value TEXT,
            new_value TEXT,
            actor TEXT,
            expected_evidence_json TEXT,
            status TEXT,
            evidence_state TEXT,
            best_score REAL,
            best_evidence_path TEXT,
            best_evidence_name TEXT,
            row_json TEXT
        );

        CREATE TABLE regulatory_requirements (
            id INTEGER PRIMARY KEY,
            categoria TEXT,
            subcategoria TEXT,
            requisito TEXT,
            periodicidade TEXT,
            prazo TEXT,
            row_json TEXT
        );

        CREATE TABLE change_proposals (
            id TEXT PRIMARY KEY,
            status TEXT,
            kind TEXT,
            domain TEXT,
            title TEXT,
            target_type TEXT,
            target_id TEXT,
            field TEXT,
            current_value TEXT,
            proposed_value TEXT,
            confidence TEXT,
            risk TEXT,
            source_type TEXT,
            source_path TEXT,
            evidence_state TEXT,
            recommended_action TEXT,
            requires_approval INTEGER,
            authorized_by TEXT,
            authorized_at TEXT,
            row_json TEXT
        );

        CREATE TABLE operational_calendar (
            date TEXT PRIMARY KEY,
            loaded INTEGER,
            status TEXT,
            points INTEGER,
            xml_families_json TEXT,
            package_families_json TEXT,
            missing_xml_families_json TEXT,
            raw_pending INTEGER,
            anp_pending INTEGER,
            alert_count INTEGER,
            pending_count INTEGER,
            open_pending_count INTEGER,
            resolved_pending_count INTEGER,
            row_json TEXT
        );

        CREATE TABLE calendar_pendencies (
            id TEXT PRIMARY KEY,
            date TEXT,
            status TEXT,
            type TEXT,
            severity TEXT,
            title TEXT,
            detail TEXT,
            recommended_action TEXT,
            resolution_mode TEXT,
            closed_by TEXT,
            closed_at TEXT,
            row_json TEXT
        );

        CREATE INDEX idx_comparisons_date_tag ON comparisons(date, tag);
        CREATE INDEX idx_alerts_severity_date ON alerts(severity, date);
        CREATE INDEX idx_events_status_timestamp ON event_evidence_events(status, timestamp);
        CREATE INDEX idx_uncertainty_date_tag ON uncertainty_monitor(date, tag);
        CREATE INDEX idx_change_proposals_status_risk ON change_proposals(status, risk);
        CREATE INDEX idx_calendar_status_date ON operational_calendar(status, date);
        CREATE INDEX idx_calendar_pendencies_status_date ON calendar_pendencies(status, date);
        """
    )

    for key, value in data.get("meta", {}).items():
        cur.execute("INSERT INTO metadata (key, value) VALUES (?, ?)", (key, json_text(value) if isinstance(value, (dict, list)) else str(value)))

    for key, value in data.get("kpis", {}).items():
        cur.execute("INSERT INTO kpis (metric, value) VALUES (?, ?)", (key, str(value)))

    for source in data.get("config", {}).get("sources", []):
        cur.execute(
            """
            INSERT INTO sources (id, label, kind, recursive, description, paths_json)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                source.get("id"),
                source.get("label"),
                source.get("kind"),
                bool_int(source.get("recursive")),
                source.get("description"),
                json_text(source.get("paths", [])),
            ),
        )

    for index, row in enumerate(data.get("files", []), start=1):
        cur.execute(
            """
            INSERT INTO files (id, date, family, family_name, kind, path, records, row_json)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (index, row.get("date"), row.get("family"), row.get("familyName"), row.get("kind"), row.get("path"), row.get("records"), json_text(row)),
        )

    for index, row in enumerate(data.get("comparisons", []), start=1):
        cur.execute(
            """
            INSERT INTO comparisons (
                id, date, tag, family, family_name, fluid, raw_corrigido, xml_corrigido,
                anp_corrigido, raw_ok, anp_ok, status, note, raw_source, xml_source, row_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                index,
                row.get("date"),
                row.get("tag"),
                row.get("family"),
                row.get("familyName"),
                row.get("fluid"),
                row.get("rawCorrigido"),
                row.get("xmlCorrigido"),
                row.get("anpCorrigido"),
                bool_int(row.get("rawOk")),
                bool_int(row.get("anpOk")),
                row.get("status"),
                row.get("note"),
                row.get("rawSource"),
                row.get("xmlSource"),
                json_text(row),
            ),
        )

    for index, row in enumerate(data.get("closing", []), start=1):
        cur.execute(
            """
            INSERT INTO closing (
                id, date, status, points, ok, warn, critical, total_oil, total_gas,
                max_bsw, missing_families_json, row_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                index,
                row.get("date"),
                row.get("status"),
                row.get("points"),
                row.get("ok"),
                row.get("warn"),
                row.get("critical"),
                row.get("totalOil"),
                row.get("totalGas"),
                row.get("maxBsw"),
                json_text(row.get("missingFamilies", [])),
                json_text(row),
            ),
        )

    for index, row in enumerate(data.get("alerts", []), start=1):
        cur.execute(
            """
            INSERT INTO alerts (id, severity, date, title, detail, area, row_json)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (index, row.get("severity"), row.get("date"), row.get("title"), row.get("detail"), row.get("area"), json_text(row)),
        )

    for index, row in enumerate(data.get("limitMonitors", []), start=1):
        pam = row.get("pam") if isinstance(row.get("pam"), dict) else {}
        cur.execute(
            """
            INSERT INTO limit_monitors (
                id, tag, family_name, fluid, current_value, min_operacao,
                max_operacao, pam, status, row_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                index,
                row.get("tag"),
                row.get("familyName"),
                row.get("fluid"),
                pam.get("value") if pam else row.get("value"),
                pam.get("lower") if pam else row.get("minOperacao"),
                pam.get("upper") if pam else row.get("maxOperacao"),
                pam.get("upper") if pam else row.get("pam"),
                row.get("status"),
                json_text(row),
            ),
        )

    for index, row in enumerate(data.get("uncertaintyMonitor", []), start=1):
        cur.execute(
            """
            INSERT INTO uncertainty_monitor (
                id, date, tag, family, family_name, uncertainty_max,
                daily_uncertainty, coverage, status, source, row_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                index,
                row.get("date"),
                row.get("tag"),
                row.get("family"),
                row.get("familyName"),
                row.get("uncertaintyMax"),
                row.get("dailyUncertainty"),
                row.get("coverage"),
                row.get("status"),
                row.get("source"),
                json_text(row),
            ),
        )

    for index, row in enumerate(data.get("eventEvidenceRadar", {}).get("events", []), start=1):
        best = (row.get("evidenceCandidates") or [{}])[0] or {}
        cur.execute(
            """
            INSERT INTO event_evidence_events (
                id, timestamp, system, flow_computer, tags_json, parameter, old_value,
                new_value, actor, expected_evidence_json, status, evidence_state,
                best_score, best_evidence_path, best_evidence_name, row_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                index,
                row.get("timestamp"),
                row.get("system"),
                row.get("flowComputer"),
                json_text(row.get("tags", [])),
                row.get("parameter"),
                row.get("oldValue"),
                row.get("newValue"),
                row.get("actor"),
                json_text(row.get("expectedEvidenceLabels", [])),
                row.get("status"),
                row.get("evidenceState"),
                row.get("bestScore"),
                best.get("path"),
                best.get("name"),
                json_text(row),
            ),
        )

    for index, row in enumerate(data.get("regulatoryMatrix", {}).get("rows", []), start=1):
        cur.execute(
            """
            INSERT INTO regulatory_requirements (
                id, categoria, subcategoria, requisito, periodicidade, prazo, row_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                index,
                row.get("Categoria"),
                row.get("Subcategoria"),
                row.get("Requisito") or row.get("Item") or row.get("Obrigacao"),
                row.get("Periodicidade"),
                row.get("Prazo"),
                json_text(row),
            ),
        )

    for row in data.get("changeProposals", []):
        cur.execute(
            """
            INSERT INTO change_proposals (
                id, status, kind, domain, title, target_type, target_id, field,
                current_value, proposed_value, confidence, risk, source_type,
                source_path, evidence_state, recommended_action, requires_approval,
                authorized_by, authorized_at, row_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                row.get("id"),
                row.get("status"),
                row.get("kind"),
                row.get("domain"),
                row.get("title"),
                row.get("targetType"),
                row.get("targetId"),
                row.get("field"),
                json_text(row.get("currentValue")),
                json_text(row.get("proposedValue")),
                row.get("confidence"),
                row.get("risk"),
                row.get("sourceType"),
                row.get("sourcePath"),
                row.get("evidenceState"),
                row.get("recommendedAction"),
                bool_int(row.get("requiresApproval")),
                row.get("authorizedBy"),
                row.get("authorizedAt"),
                json_text(row),
            ),
        )

    for row in data.get("operationalCalendar", {}).get("days", []):
        cur.execute(
            """
            INSERT INTO operational_calendar (
                date, loaded, status, points, xml_families_json, package_families_json,
                missing_xml_families_json, raw_pending, anp_pending, alert_count,
                pending_count, open_pending_count, resolved_pending_count, row_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                row.get("date"),
                bool_int(row.get("loaded")),
                row.get("status"),
                row.get("points"),
                json_text(row.get("xmlFamilies", [])),
                json_text(row.get("packageFamilies", [])),
                json_text(row.get("missingXmlFamilies", [])),
                row.get("rawPending"),
                row.get("anpPending"),
                row.get("alertCount"),
                row.get("pendingCount"),
                row.get("openPendingCount"),
                row.get("resolvedPendingCount"),
                json_text(row),
            ),
        )
        for item in row.get("pendingItems", []):
            cur.execute(
                """
                INSERT INTO calendar_pendencies (
                    id, date, status, type, severity, title, detail, recommended_action,
                    resolution_mode, closed_by, closed_at, row_json
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    item.get("id"),
                    row.get("date"),
                    item.get("status"),
                    item.get("type"),
                    item.get("severity"),
                    item.get("title"),
                    item.get("detail"),
                    item.get("recommendedAction"),
                    item.get("resolutionMode"),
                    item.get("closedBy"),
                    item.get("closedAt"),
                    json_text(item),
                ),
            )

    conn.commit()

    table_names = [
        "metadata",
        "kpis",
        "sources",
        "files",
        "comparisons",
        "closing",
        "alerts",
        "limit_monitors",
        "uncertainty_monitor",
        "event_evidence_events",
        "regulatory_requirements",
        "change_proposals",
        "operational_calendar",
        "calendar_pendencies",
    ]
    counts = {name: cur.execute(f"SELECT COUNT(*) FROM {name}").fetchone()[0] for name in table_names}
    conn.close()

    return {
        "path": str(SQLITE_PATH),
        "generatedAt": data.get("meta", {}).get("generatedAt"),
        "tables": table_names,
        "tableCounts": counts,
        "sizeBytes": SQLITE_PATH.stat().st_size if SQLITE_PATH.exists() else 0,
    }


def main() -> None:
    xml_records, files = parse_measurement_xmls()
    cv_index = parse_cv_daily()
    anp_index, anp_rows, latest_points = parse_anp_exports()
    points = parse_points()
    bsw_rows, max_bsw = parse_bsw()
    failures = parse_failures()
    operator_panel_health = build_operator_panel_health()
    mpfm = parse_mpfm()
    lab_report = parse_lab_report()
    measurement_models = parse_measurement_models()
    regulatory_matrix = parse_regulatory_matrix()
    event_evidence_radar = build_event_evidence_radar()
    ihm_reports = parse_ihm_daily_reports()
    gas_balance = parse_gas_balance()
    cartas_anp = parse_cartas_anp()
    calibration_control = parse_calibration_control()
    operating_ranges = parse_operating_ranges()
    sfp_registration = parse_sfp_registration()
    alarm_management = parse_alarm_management()
    comparisons = compare_layers(xml_records, anp_index, cv_index)
    closing = build_daily_closing(comparisons, files, bsw_rows)
    alerts = build_alerts(closing, comparisons, failures, mpfm, operator_panel_health)

    dates = sorted({row["date"] for row in comparisons if row.get("date")})
    latest_anp_date = max((row["date"] for row in anp_rows), default=None)
    spec_points = []
    for row in latest_points:
        tag = row["tag"]
        point = points.get(tag, {})
        value = row.get("volumeCorrigido") or 0
        max_op = point.get("maxOperacao")
        min_op = point.get("minOperacao")
        in_range = True
        if max_op is not None and value > max_op:
            in_range = False
        if min_op is not None and value < min_op:
            in_range = False
        spec_points.append({**row, **point, "inRange": in_range})
    limit_monitors = build_limit_monitors(spec_points, points)
    uncertainty_monitor = build_uncertainty_monitor(comparisons, points)
    ai_modules = build_ai_modules(alerts)
    change_proposals = build_change_proposals(limit_monitors, uncertainty_monitor, regulatory_matrix, event_evidence_radar)
    operational_calendar = build_operational_calendar(dates, files, comparisons, closing, alerts)

    total_xml_anp = len(comparisons)
    ok_xml_anp = sum(1 for row in comparisons if row["anpOk"])
    ok_raw = sum(1 for row in comparisons if row["rawOk"])
    data = {
        "meta": {
            "generatedAt": datetime.now().isoformat(timespec="seconds"),
            "workspace": str(ROOT),
            "configPath": str(CONFIG_PATH),
            "sourceDates": dates,
            "latestAnpDate": latest_anp_date,
            "today": TODAY.isoformat(),
        },
        "config": CONFIG,
        "kpis": {
            "xmlFiles": len(files),
            "xmlRecords": len(xml_records),
            "anpRows": len(anp_rows),
            "comparisonRows": total_xml_anp,
            "xmlAnpOk": ok_xml_anp,
            "rawXmlOk": ok_raw,
            "openFailures": failures.get("open", 0),
            "mpfmAlerts": len(mpfm.get("alerts", [])),
            "modelFiles": measurement_models["summary"].get("files", 0),
            "modelSignals": measurement_models["summary"].get("signals", 0),
            "operatorPanelReady": operator_panel_health.get("ready", 0),
            "operatorPanelMissing": (operator_panel_health.get("missingFiles", 0) or 0) + (operator_panel_health.get("missingInformation", 0) or 0),
            "eventEvidenceCritical": event_evidence_radar["summary"].get("critical", 0),
        },
        "families": FAMILIES,
        "files": files,
        "closing": closing,
        "comparisons": comparisons,
        "latestPoints": spec_points,
        "limitMonitors": limit_monitors,
        "uncertaintyMonitor": uncertainty_monitor,
        "analytical": {
            "labReport": lab_report,
            "bsw": {"rows": bsw_rows, "max": max_bsw},
        },
        "measurementModels": measurement_models,
        "operatorPanelHealth": operator_panel_health,
        "ai": ai_modules,
        "regulatoryMatrix": regulatory_matrix,
        "eventEvidenceRadar": event_evidence_radar,
        "changeProposals": change_proposals,
        "operationalCalendar": operational_calendar,
        "bsw": {"rows": bsw_rows, "max": max_bsw},
        "failures": failures,
        "mpfm": mpfm,
        "alerts": alerts,
        "ihmReports": ihm_reports,
        "gasBalance": gas_balance,
        "cartasAnp": cartas_anp,
        "calibrationControl": calibration_control,
        "operatingRanges": operating_ranges,
        "sfpRegistration": sfp_registration,
        "alarmManagement": alarm_management,
    }
    data = clean_json(data)
    database_summary = write_sqlite_database(data)
    data["database"] = database_summary
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT.write_text(json.dumps(data, ensure_ascii=False, indent=2, allow_nan=False), encoding="utf-8")
    print(f"wrote {OUTPUT}")
    print(f"wrote {SQLITE_PATH}")
    print(
        json.dumps(
            {
                "dates": dates,
                "comparisons": total_xml_anp,
                "xmlAnpOk": ok_xml_anp,
                "rawXmlOk": ok_raw,
                "alerts": len(alerts),
                "proposals": len(change_proposals),
                "measurementModels": measurement_models["summary"],
                "calendar": operational_calendar["summary"],
                "eventEvidence": event_evidence_radar["summary"],
                "sqlite": database_summary["tableCounts"],
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
